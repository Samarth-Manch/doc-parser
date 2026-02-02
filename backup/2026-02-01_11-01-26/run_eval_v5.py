#!/usr/bin/env python3
"""
Comprehensive Rule Extraction Evaluation Script v2
Evaluates generated schema against BUD document and human-made reference.

Handles the formFillMetadatas flat structure with actionType rules.
"""

import json
import os
import re
from collections import defaultdict
from datetime import datetime
from typing import Dict, List, Any, Set, Tuple
from docx import Document

# File paths
GENERATED_FILE = "/home/samart/project/doc-parser/adws/2026-02-01_11-01-26/populated_schema_v5.json"
REFERENCE_FILE = "/home/samart/project/doc-parser/documents/json_output/vendor_creation_sample_bud.json"
BUD_FILE = "/home/samart/project/doc-parser/documents/Vendor Creation Sample BUD.docx"
OUTPUT_FILE = "/home/samart/project/doc-parser/adws/2026-02-01_11-01-26/eval_report_v5.json"

# Configuration
PASS_THRESHOLD = 0.9
ITERATION = 5


def load_json(file_path: str) -> Dict:
    """Load JSON file."""
    with open(file_path, 'r', encoding='utf-8') as f:
        return json.load(f)


def parse_bud_document(file_path: str) -> Dict[str, List[str]]:
    """Parse BUD document to extract natural language rules by panel."""
    doc = Document(file_path)

    panel_rules = defaultdict(list)
    current_panel = "General"
    panel_pattern = re.compile(r'panel\s*\d*:?\s*(.+)', re.IGNORECASE)

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Detect panel headers
        lower_text = text.lower()
        panel_match = panel_pattern.search(text)
        if panel_match or 'panel' in lower_text:
            if panel_match:
                current_panel = panel_match.group(1).strip()
            else:
                current_panel = text

        # Detect rule-related content
        if any(keyword in lower_text for keyword in ['if', 'when', 'must', 'should', 'visible', 'mandatory', 'disabled', 'verify', 'ocr', 'dropdown', 'validation']):
            panel_rules[current_panel].append(text)

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            for cell_text in cells:
                if len(cell_text) > 10:  # Skip short cells
                    lower_cell = cell_text.lower()
                    if any(keyword in lower_cell for keyword in ['if', 'when', 'must', 'visible', 'mandatory', 'disabled', 'verify', 'logic', 'rule', 'dropdown']):
                        panel_rules[current_panel].append(cell_text)

    return dict(panel_rules)


def extract_form_fill_metadatas(schema: Dict) -> List[Dict]:
    """Extract formFillMetadatas from schema."""
    return schema.get('template', {}).get('documentTypes', [{}])[0].get('formFillMetadatas', [])


def get_field_name(field: Dict) -> str:
    """Get field name/identifier."""
    return field.get('variableName') or field.get('formTag') or 'unknown'


def group_fields_by_page(fields: List[Dict]) -> Dict[int, List[Dict]]:
    """Group fields by page number (as proxy for panels)."""
    by_page = defaultdict(list)
    for field in fields:
        page = field.get('page', 1)
        by_page[page].append(field)
    return dict(by_page)


def extract_rules_by_type(fields: List[Dict]) -> Dict[str, List[Dict]]:
    """Extract all rules organized by action type."""
    rules_by_type = defaultdict(list)

    for field in fields:
        field_name = get_field_name(field)
        form_fill_rules = field.get('formFillRules', [])

        for rule in form_fill_rules:
            action_type = rule.get('actionType', 'UNKNOWN')
            rule_info = {
                'fieldName': field_name,
                'fieldId': field.get('id'),
                'rule': rule,
                'actionType': action_type
            }
            rules_by_type[action_type].append(rule_info)

    return dict(rules_by_type)


def count_rules_by_type(fields: List[Dict]) -> Dict[str, int]:
    """Count rules by action type."""
    counts = defaultdict(int)
    for field in fields:
        for rule in field.get('formFillRules', []):
            action_type = rule.get('actionType', 'UNKNOWN')
            counts[action_type] += 1
    return dict(counts)


def check_ocr_chains(gen_rules_by_type: Dict) -> Dict:
    """Check OCR rules have proper postTriggerRuleIds chains to VERIFY."""
    results = {
        'ocr_rules_count': 0,
        'ocr_with_chains': [],
        'ocr_missing_chains': [],
        'issues': []
    }

    for rule_info in gen_rules_by_type.get('OCR', []):
        results['ocr_rules_count'] += 1
        rule = rule_info['rule']
        field_name = rule_info['fieldName']
        post_trigger = rule.get('postTriggerRuleIds', [])

        if post_trigger:
            results['ocr_with_chains'].append({
                'field': field_name,
                'fieldId': rule_info['fieldId'],
                'postTriggerRuleIds': post_trigger
            })
        else:
            results['ocr_missing_chains'].append({
                'field': field_name,
                'fieldId': rule_info['fieldId']
            })
            results['issues'].append(f"OCR rule for '{field_name}' missing postTriggerRuleIds chain")

    return results


def check_verify_ordinal(gen_rules_by_type: Dict) -> Dict:
    """Check VERIFY rules have proper ordinal destinationIds mapping."""
    results = {
        'verify_rules_count': 0,
        'verify_with_ordinal': [],
        'verify_missing_ordinal': [],
        'issues': []
    }

    for rule_info in gen_rules_by_type.get('VERIFY', []):
        results['verify_rules_count'] += 1
        rule = rule_info['rule']
        field_name = rule_info['fieldName']
        destination_ids = rule.get('destinationIds', [])

        if destination_ids:
            results['verify_with_ordinal'].append({
                'field': field_name,
                'fieldId': rule_info['fieldId'],
                'destinationIds': destination_ids
            })
        else:
            results['verify_missing_ordinal'].append({
                'field': field_name,
                'fieldId': rule_info['fieldId']
            })
            results['issues'].append(f"VERIFY rule for '{field_name}' missing destinationIds")

    return results


def check_visibility_pairs(gen_rules_by_type: Dict) -> Dict:
    """Check MAKE_VISIBLE and MAKE_INVISIBLE pairs exist."""
    results = {
        'visible_count': 0,
        'invisible_count': 0,
        'visible_fields': [],
        'invisible_fields': [],
        'unpaired_visible': [],
        'unpaired_invisible': [],
        'issues': []
    }

    visible_field_set = set()
    invisible_field_set = set()

    for rule_info in gen_rules_by_type.get('MAKE_VISIBLE', []):
        results['visible_count'] += 1
        field_name = rule_info['fieldName']
        visible_field_set.add(field_name)
        results['visible_fields'].append(field_name)

    for rule_info in gen_rules_by_type.get('MAKE_INVISIBLE', []):
        results['invisible_count'] += 1
        field_name = rule_info['fieldName']
        invisible_field_set.add(field_name)
        results['invisible_fields'].append(field_name)

    # Check for unpaired rules
    for field in visible_field_set - invisible_field_set:
        results['unpaired_visible'].append(field)
        results['issues'].append(f"MAKE_VISIBLE for '{field}' has no corresponding MAKE_INVISIBLE")

    for field in invisible_field_set - visible_field_set:
        results['unpaired_invisible'].append(field)

    return results


def check_mandatory_pairs(gen_rules_by_type: Dict) -> Dict:
    """Check MAKE_MANDATORY and MAKE_NON_MANDATORY pairs."""
    results = {
        'mandatory_count': 0,
        'non_mandatory_count': 0,
        'mandatory_fields': [],
        'non_mandatory_fields': [],
        'issues': []
    }

    for rule_info in gen_rules_by_type.get('MAKE_MANDATORY', []):
        results['mandatory_count'] += 1
        results['mandatory_fields'].append(rule_info['fieldName'])

    for rule_info in gen_rules_by_type.get('MAKE_NON_MANDATORY', []):
        results['non_mandatory_count'] += 1
        results['non_mandatory_fields'].append(rule_info['fieldName'])

    return results


def check_ext_dropdown_pairs(gen_rules_by_type: Dict) -> Dict:
    """Check EXT_DROP_DOWN has corresponding EXT_VALUE rules."""
    results = {
        'ext_dropdown_count': 0,
        'ext_value_count': 0,
        'ext_dropdown_fields': [],
        'ext_value_fields': [],
        'issues': []
    }

    for rule_info in gen_rules_by_type.get('EXT_DROP_DOWN', []):
        results['ext_dropdown_count'] += 1
        results['ext_dropdown_fields'].append(rule_info['fieldName'])

    for rule_info in gen_rules_by_type.get('EXT_VALUE', []):
        results['ext_value_count'] += 1
        results['ext_value_fields'].append(rule_info['fieldName'])

    return results


def check_disabled_consolidated(gen_rules_by_type: Dict) -> Dict:
    """Check MAKE_DISABLED rules are consolidated (not individual)."""
    results = {
        'disabled_count': 0,
        'disabled_rules': [],
        'avg_destinations_per_rule': 0,
        'issues': []
    }

    total_destinations = 0
    for rule_info in gen_rules_by_type.get('MAKE_DISABLED', []):
        results['disabled_count'] += 1
        rule = rule_info['rule']
        dest_ids = rule.get('destinationIds', [])
        total_destinations += len(dest_ids)
        results['disabled_rules'].append({
            'field': rule_info['fieldName'],
            'destination_count': len(dest_ids)
        })

    if results['disabled_count'] > 0:
        results['avg_destinations_per_rule'] = round(total_destinations / results['disabled_count'], 2)
        # Flag if rules are too fragmented (less than 5 destinations average suggests non-consolidated)
        if results['avg_destinations_per_rule'] < 5 and results['disabled_count'] > 3:
            results['issues'].append("MAKE_DISABLED rules may not be consolidated (low avg destinations)")

    return results


def compare_rule_counts(gen_counts: Dict, ref_counts: Dict) -> Dict:
    """Compare rule type counts between generated and reference."""
    all_types = set(gen_counts.keys()) | set(ref_counts.keys())

    comparison = {}
    for rule_type in sorted(all_types):
        gen = gen_counts.get(rule_type, 0)
        ref = ref_counts.get(rule_type, 0)
        comparison[rule_type] = {
            'generated': gen,
            'reference': ref,
            'difference': gen - ref,
            'match': gen == ref
        }

    return comparison


def find_missing_rule_types(gen_counts: Dict, ref_counts: Dict) -> List[Dict]:
    """Find rule types present in reference but missing or under-represented in generated."""
    missing = []

    for rule_type, ref_count in ref_counts.items():
        gen_count = gen_counts.get(rule_type, 0)
        if gen_count < ref_count:
            severity = 'high' if rule_type in ['VERIFY', 'OCR', 'MAKE_MANDATORY', 'EXT_DROP_DOWN'] else 'medium'
            missing.append({
                'rule_type': rule_type,
                'generated': gen_count,
                'reference': ref_count,
                'shortfall': ref_count - gen_count,
                'severity': severity
            })

    return sorted(missing, key=lambda x: (0 if x['severity'] == 'high' else 1, -x['shortfall']))


def calculate_overall_score(gen_counts: Dict, ref_counts: Dict, critical_checks: Dict) -> float:
    """Calculate overall evaluation score."""
    # Base score from rule count matching
    all_types = set(gen_counts.keys()) | set(ref_counts.keys())
    if not all_types:
        return 0.0

    type_scores = []
    for rule_type in all_types:
        gen = gen_counts.get(rule_type, 0)
        ref = ref_counts.get(rule_type, 0)
        if ref > 0:
            type_scores.append(min(gen / ref, 1.0))
        elif gen == 0:
            type_scores.append(1.0)  # Both zero is fine
        else:
            type_scores.append(0.8)  # Extra rules generated - slightly penalize

    base_score = sum(type_scores) / len(type_scores) if type_scores else 0.0

    # Penalty for critical issues
    penalties = 0.0

    # OCR chains penalty
    ocr_check = critical_checks.get('ocr_chains', {})
    if ocr_check.get('ocr_missing_chains'):
        ocr_total = ocr_check.get('ocr_rules_count', 1)
        ocr_missing = len(ocr_check.get('ocr_missing_chains', []))
        penalties += 0.1 * (ocr_missing / ocr_total) if ocr_total > 0 else 0

    # VERIFY ordinal penalty
    verify_check = critical_checks.get('ordinal_mapping', {})
    if verify_check.get('verify_missing_ordinal'):
        verify_total = verify_check.get('verify_rules_count', 1)
        verify_missing = len(verify_check.get('verify_missing_ordinal', []))
        penalties += 0.1 * (verify_missing / verify_total) if verify_total > 0 else 0

    # Visibility pairs penalty
    vis_check = critical_checks.get('visibility_pairs', {})
    if vis_check.get('unpaired_visible'):
        penalties += 0.05 * min(len(vis_check.get('unpaired_visible', [])) / 10, 1.0)

    final_score = max(0.0, base_score - penalties)
    return round(final_score, 3)


def generate_self_heal_instructions(eval_results: Dict) -> Dict:
    """Generate self-heal instructions based on evaluation."""
    priority_fixes = []

    # Check OCR chain issues
    ocr_check = eval_results['critical_checks']['ocr_chains']
    if ocr_check.get('ocr_missing_chains'):
        priority_fixes.append({
            'priority': 1,
            'category': 'OCR_CHAINS',
            'description': 'Add postTriggerRuleIds to OCR rules to chain to VERIFY rules',
            'affected_fields': [item['field'] for item in ocr_check['ocr_missing_chains']],
            'action': 'For each OCR rule, add postTriggerRuleIds array containing the ID of the corresponding VERIFY rule'
        })

    # Check ordinal mapping issues
    verify_check = eval_results['critical_checks']['ordinal_mapping']
    if verify_check.get('verify_missing_ordinal'):
        priority_fixes.append({
            'priority': 2,
            'category': 'VERIFY_ORDINAL',
            'description': 'Add destinationIds to VERIFY rules for ordinal field mapping',
            'affected_fields': [item['field'] for item in verify_check['verify_missing_ordinal']],
            'action': 'For each VERIFY rule, add destinationIds array containing IDs of fields that should receive the verified value'
        })

    # Check visibility pair issues
    vis_check = eval_results['critical_checks']['visibility_pairs']
    if vis_check.get('unpaired_visible'):
        priority_fixes.append({
            'priority': 3,
            'category': 'VISIBILITY_PAIRS',
            'description': 'Add MAKE_INVISIBLE rules for unpaired MAKE_VISIBLE rules',
            'affected_fields': vis_check['unpaired_visible'],
            'action': 'For each MAKE_VISIBLE rule, create a corresponding MAKE_INVISIBLE rule with opposite condition'
        })

    # Add missing rule types from comparison
    for item in eval_results.get('missing_rule_types', []):
        if item['shortfall'] > 0:
            priority = 4 if item['severity'] == 'high' else 5
            priority_fixes.append({
                'priority': priority,
                'category': f'MISSING_{item["rule_type"]}',
                'description': f'Add {item["shortfall"]} more {item["rule_type"]} rules',
                'shortfall': item['shortfall'],
                'action': f'Generate {item["shortfall"]} additional {item["rule_type"]} rules based on BUD document logic'
            })

    # Sort by priority
    priority_fixes.sort(key=lambda x: x['priority'])

    return {
        'priority_fixes': priority_fixes,
        'total_issues': len(priority_fixes),
        'critical_issues': len([f for f in priority_fixes if f['priority'] <= 2])
    }


def evaluate_by_page(gen_fields: List[Dict], ref_fields: List[Dict]) -> List[Dict]:
    """Evaluate rules page by page (as proxy for panels)."""
    gen_by_page = group_fields_by_page(gen_fields)
    ref_by_page = group_fields_by_page(ref_fields)

    all_pages = sorted(set(gen_by_page.keys()) | set(ref_by_page.keys()))

    page_results = []
    for page in all_pages:
        gen_page_fields = gen_by_page.get(page, [])
        ref_page_fields = ref_by_page.get(page, [])

        gen_rule_count = sum(len(f.get('formFillRules', [])) for f in gen_page_fields)
        ref_rule_count = sum(len(f.get('formFillRules', [])) for f in ref_page_fields)

        gen_counts = count_rules_by_type(gen_page_fields)
        ref_counts = count_rules_by_type(ref_page_fields)

        # Calculate page score
        if ref_rule_count > 0:
            coverage = min(gen_rule_count / ref_rule_count, 1.0)
        else:
            coverage = 1.0 if gen_rule_count == 0 else 0.8

        page_results.append({
            'page': page,
            'generated_field_count': len(gen_page_fields),
            'reference_field_count': len(ref_page_fields),
            'generated_rule_count': gen_rule_count,
            'reference_rule_count': ref_rule_count,
            'coverage': round(coverage, 3),
            'generated_rule_types': gen_counts,
            'reference_rule_types': ref_counts,
            'passed': coverage >= PASS_THRESHOLD
        })

    return page_results


def main():
    print("=" * 70)
    print("COMPREHENSIVE RULE EXTRACTION EVALUATION v2")
    print(f"Iteration: {ITERATION}")
    print(f"Pass Threshold: {PASS_THRESHOLD}")
    print("=" * 70)

    # Load data
    print("\n[1/7] Loading generated schema...")
    generated = load_json(GENERATED_FILE)
    gen_fields = extract_form_fill_metadatas(generated)
    print(f"   Fields: {len(gen_fields)}")

    print("[2/7] Loading reference schema...")
    reference = load_json(REFERENCE_FILE)
    ref_fields = extract_form_fill_metadatas(reference)
    print(f"   Fields: {len(ref_fields)}")

    print("[3/7] Parsing BUD document...")
    bud_rules = parse_bud_document(BUD_FILE)
    print(f"   Panels with rules found: {len(bud_rules)}")

    # Count rules
    print("[4/7] Counting rules by type...")
    gen_counts = count_rules_by_type(gen_fields)
    ref_counts = count_rules_by_type(ref_fields)
    print(f"   Generated total rules: {sum(gen_counts.values())}")
    print(f"   Reference total rules: {sum(ref_counts.values())}")

    # Extract rules by type for detailed checks
    print("[5/7] Extracting rules for analysis...")
    gen_rules_by_type = extract_rules_by_type(gen_fields)

    # Critical checks
    print("[6/7] Running critical checks...")
    ocr_chains = check_ocr_chains(gen_rules_by_type)
    ordinal_mapping = check_verify_ordinal(gen_rules_by_type)
    visibility_pairs = check_visibility_pairs(gen_rules_by_type)
    mandatory_pairs = check_mandatory_pairs(gen_rules_by_type)
    ext_dropdown = check_ext_dropdown_pairs(gen_rules_by_type)
    disabled_check = check_disabled_consolidated(gen_rules_by_type)

    critical_checks = {
        'ocr_chains': ocr_chains,
        'ordinal_mapping': ordinal_mapping,
        'visibility_pairs': visibility_pairs,
        'mandatory_pairs': mandatory_pairs,
        'ext_dropdown_pairs': ext_dropdown,
        'disabled_consolidated': disabled_check
    }

    # Page-by-page evaluation
    print("[7/7] Evaluating by page...")
    page_results = evaluate_by_page(gen_fields, ref_fields)

    # Rule type comparison
    rule_comparison = compare_rule_counts(gen_counts, ref_counts)

    # Missing rules
    missing_rule_types = find_missing_rule_types(gen_counts, ref_counts)

    # Calculate overall score
    overall_score = calculate_overall_score(gen_counts, ref_counts, critical_checks)

    # Build eval report
    eval_report = {
        'metadata': {
            'generated_file': GENERATED_FILE,
            'reference_file': REFERENCE_FILE,
            'bud_file': BUD_FILE,
            'iteration': ITERATION,
            'pass_threshold': PASS_THRESHOLD,
            'evaluation_timestamp': datetime.now().isoformat()
        },
        'summary': {
            'overall_score': overall_score,
            'passed': overall_score >= PASS_THRESHOLD,
            'generated_field_count': len(gen_fields),
            'reference_field_count': len(ref_fields),
            'generated_rule_count': sum(gen_counts.values()),
            'reference_rule_count': sum(ref_counts.values()),
            'pages_evaluated': len(page_results),
            'pages_passed': len([p for p in page_results if p['passed']]),
            'pages_failed': len([p for p in page_results if not p['passed']])
        },
        'page_scores': page_results,
        'rule_type_comparison': rule_comparison,
        'critical_checks': critical_checks,
        'missing_rule_types': missing_rule_types,
        'bud_natural_language_rules': bud_rules,
        'self_heal_instructions': {}
    }

    # Generate self-heal instructions
    eval_report['self_heal_instructions'] = generate_self_heal_instructions(eval_report)

    # Save report
    with open(OUTPUT_FILE, 'w', encoding='utf-8') as f:
        json.dump(eval_report, f, indent=2, ensure_ascii=False)

    # Print summary
    print("\n" + "=" * 70)
    print("EVALUATION SUMMARY")
    print("=" * 70)
    print(f"Overall Score: {overall_score:.1%}")
    print(f"Status: {'PASSED ✓' if overall_score >= PASS_THRESHOLD else 'FAILED ✗'}")
    print(f"\nFields: Generated={len(gen_fields)}, Reference={len(ref_fields)}")
    print(f"Rules:  Generated={sum(gen_counts.values())}, Reference={sum(ref_counts.values())}")
    print(f"Pages:  Evaluated={len(page_results)}, Passed={len([p for p in page_results if p['passed']])}")
    print(f"\nCritical Issues: {eval_report['self_heal_instructions']['critical_issues']}")
    print(f"Total Issues: {eval_report['self_heal_instructions']['total_issues']}")

    # Print rule type comparison
    print("\n" + "-" * 50)
    print("RULE TYPE COMPARISON")
    print("-" * 50)
    print(f"{'Rule Type':<30} {'Gen':>6} {'Ref':>6} {'Diff':>6} {'Match'}")
    print("-" * 50)
    for rule_type, counts in sorted(rule_comparison.items(), key=lambda x: -(x[1]['reference'])):
        diff = counts['difference']
        diff_str = f"+{diff}" if diff > 0 else str(diff)
        match_mark = "✓" if counts['match'] else "✗"
        print(f"{rule_type:<30} {counts['generated']:>6} {counts['reference']:>6} {diff_str:>6} {match_mark}")

    # Print critical check summaries
    print("\n" + "-" * 50)
    print("CRITICAL CHECKS")
    print("-" * 50)
    print(f"OCR Chains: {ocr_chains['ocr_rules_count']} OCR rules, {len(ocr_chains['ocr_with_chains'])} with chains, {len(ocr_chains['ocr_missing_chains'])} missing")
    print(f"VERIFY Ordinal: {ordinal_mapping['verify_rules_count']} VERIFY rules, {len(ordinal_mapping['verify_with_ordinal'])} with mapping")
    print(f"Visibility Pairs: {visibility_pairs['visible_count']} visible, {visibility_pairs['invisible_count']} invisible, {len(visibility_pairs['unpaired_visible'])} unpaired")
    print(f"Mandatory Pairs: {mandatory_pairs['mandatory_count']} mandatory, {mandatory_pairs['non_mandatory_count']} non-mandatory")
    print(f"Ext Dropdown: {ext_dropdown['ext_dropdown_count']} EXT_DROP_DOWN, {ext_dropdown['ext_value_count']} EXT_VALUE")
    print(f"Disabled: {disabled_check['disabled_count']} rules, avg {disabled_check['avg_destinations_per_rule']} destinations/rule")

    # Print priority fixes
    if eval_report['self_heal_instructions']['priority_fixes']:
        print("\n" + "-" * 50)
        print("PRIORITY FIXES NEEDED")
        print("-" * 50)
        for fix in eval_report['self_heal_instructions']['priority_fixes'][:7]:
            affected = fix.get('affected_fields', [])
            affected_str = f"({len(affected)} fields)" if affected else f"(shortfall: {fix.get('shortfall', 0)})"
            print(f"[P{fix['priority']}] {fix['category']}: {fix['description']} {affected_str}")

    print(f"\nReport saved to: {OUTPUT_FILE}")
    return eval_report


if __name__ == "__main__":
    main()
