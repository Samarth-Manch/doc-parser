"""
Core document parser for OOXML files.
"""

import re
import zipfile
import base64
from pathlib import Path
from typing import Optional
from docx import Document
from docx.table import Table
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .models import (
    ParsedDocument,
    DocumentMetadata,
    VersionEntry,
    Section,
    TableData,
    FieldDefinition,
    FieldType,
    WorkflowStep,
    ApprovalRule,
    IntegrationField,
    DocumentRequirementMatrix,
    ImageReference,
    FontInfo,
    ParagraphFormatting,
    RunFormatting,
    CellFormatting,
    TableFormatting,
    DocumentElement,
    PageSetup,
    HeaderFooter,
)


class DocumentParser:
    """
    Parser for OOXML documents that extracts fields, rules, workflows, and metadata.

    Optimized for the Vendor Creation template format but handles variations
    in other document structures.
    """

    # Patterns for identifying table types
    FIELD_HEADER_PATTERNS = [
        ["field name", "field type"],
        ["field name", "type"],
        ["filed name", "field type"],  # Handle typos
    ]

    APPROVAL_KEYWORDS = ["approver", "approval", "tse", "asm"]
    INITIATOR_KEYWORDS = ["initiator", "requestor", "requester", "first party"]
    SPOC_KEYWORDS = ["spoc", "vendor", "second party"]

    # Patterns to identify the MASTER field table (Section 4.4 Field-Level Information)
    # vs workflow/actor tables (Section 4.5.x Behaviour tables)
    MASTER_TABLE_CONTEXT_PATTERNS = [
        "4.4",
        "field-level information",
        "field level information",
        "master field",
        "all fields",
    ]

    # Patterns that indicate this is a workflow/actor table, NOT the master table
    WORKFLOW_TABLE_CONTEXT_PATTERNS = [
        "4.5",
        "behaviour",
        "behavior",
        "initiator",
        "vendor",
        "approver",
        "contributor",
        "spoc",
        "first party",
        "second party",
    ]

    def __init__(self):
        self._current_section_context = ""
        self._current_actor_context = ""

    def parse(self, file_path: str) -> ParsedDocument:
        """
        Parse a document and extract all structured information.

        Args:
            file_path: Path to the .docx file

        Returns:
            ParsedDocument with all extracted information
        """
        doc = Document(file_path)

        # Initialize result
        result = ParsedDocument(
            file_path=str(file_path),
            metadata=self._extract_metadata(doc),
        )

        # Extract visual elements
        result.images = self._extract_images(file_path, doc)

        # Extract page setup
        result.page_setup = self._extract_page_setup(doc)

        # Extract headers and footers
        result.header, result.footer = self._extract_headers_footers(doc, file_path)

        # Extract EXACT document structure (for perfect recreation)
        result.document_elements = self._extract_exact_document_order(doc, file_path)

        # Extract document structure
        result.sections = self._extract_sections(doc)
        result.raw_tables = self._extract_all_tables(doc)

        # Extract specific content types
        self._extract_version_history(doc, result)
        self._extract_terminology(doc, result)
        self._extract_fields_from_tables(doc, result)
        self._extract_workflows(doc, result)
        self._extract_approval_rules(doc, result)
        self._extract_scope(doc, result)
        self._extract_reference_tables(doc, result)
        self._extract_integration_info(doc, result)
        self._extract_document_requirements(doc, result)
        self._extract_dropdown_values(result)

        return result

    def _extract_metadata(self, doc: Document) -> DocumentMetadata:
        """Extract document core properties and metadata."""
        props = doc.core_properties

        # Try to extract process name from title or first heading
        process_name = ""
        for para in doc.paragraphs[:10]:  # Check first 10 paragraphs
            if para.style and "Heading 1" in para.style.name:
                process_name = para.text.strip()
                break

        return DocumentMetadata(
            title=props.title or "",
            author=props.author or "",
            subject=props.subject or "",
            created=str(props.created) if props.created else None,
            modified=str(props.modified) if props.modified else None,
            last_modified_by=props.last_modified_by or "",
            process_name=process_name,
        )

    def _extract_sections(self, doc: Document) -> list[Section]:
        """Extract hierarchical section structure from document."""
        sections: list[Section] = []
        section_stack: list[Section] = []
        current_content: list[str] = []
        current_runs: list[list[RunFormatting]] = []
        current_para_formats: list[ParagraphFormatting] = []

        for element in doc.element.body:
            if element.tag.endswith("p"):
                # Find corresponding paragraph
                para = None
                for p in doc.paragraphs:
                    if p._element is element:
                        para = p
                        break

                if para is None:
                    continue

                style_name = para.style.name if para.style else ""
                text = para.text.strip()

                # Check if this is a heading
                heading_level = self._get_heading_level(style_name)

                if heading_level > 0:
                    # Save current content to previous section
                    if section_stack and current_content:
                        section_stack[-1].content.extend(current_content)
                        section_stack[-1].runs.extend(current_runs)
                        section_stack[-1].paragraph_formats.extend(current_para_formats)
                        current_content = []
                        current_runs = []
                        current_para_formats = []

                    # Create new section
                    new_section = Section(
                        heading=text,
                        level=heading_level,
                        content=[],
                    )

                    # Find appropriate parent
                    while section_stack and section_stack[-1].level >= heading_level:
                        section_stack.pop()

                    if section_stack:
                        section_stack[-1].subsections.append(new_section)
                    else:
                        sections.append(new_section)

                    section_stack.append(new_section)
                    self._current_section_context = text

                elif text:
                    current_content.append(text)
                    # Extract formatting for this paragraph
                    current_runs.append(self._extract_run_formatting(para))
                    current_para_formats.append(self._extract_paragraph_formatting(para))

        # Add remaining content
        if section_stack and current_content:
            section_stack[-1].content.extend(current_content)
            section_stack[-1].runs.extend(current_runs)
            section_stack[-1].paragraph_formats.extend(current_para_formats)

        return sections

    def _get_heading_level(self, style_name: str) -> int:
        """Extract heading level from style name."""
        if not style_name:
            return 0

        if style_name == "Title":
            return 0  # Title is not a section heading

        match = re.search(r"Heading\s*(\d+)", style_name, re.IGNORECASE)
        if match:
            return int(match.group(1))

        return 0

    def _extract_all_tables(self, doc: Document) -> list[TableData]:
        """Extract all tables with their content and context."""
        tables: list[TableData] = []
        current_heading = ""

        for element in doc.element.body:
            if element.tag.endswith("p"):
                for para in doc.paragraphs:
                    if para._element is element:
                        if para.style and "Heading" in para.style.name:
                            current_heading = para.text.strip()
                        break

            elif element.tag.endswith("tbl"):
                for table in doc.tables:
                    if table._tbl is element:
                        table_data = self._parse_table(table, current_heading)
                        tables.append(table_data)
                        break

        return tables

    def _parse_table(self, table: Table, context: str = "") -> TableData:
        """Parse a table into structured TableData."""
        rows = []
        headers = []

        for i, row in enumerate(table.rows):
            cells = [self._clean_cell_text(cell.text) for cell in row.cells]

            if i == 0:
                headers = cells
            else:
                rows.append(cells)

        # Check if headers look invalid (merged cells with repeated text)
        # If so, use first data row as headers
        if headers and len(set(headers)) == 1 and rows:
            # All headers are the same - likely merged cells
            # Use first data row as headers
            headers = rows[0]
            rows = rows[1:]

        # Determine table type
        table_type = self._identify_table_type(headers, rows, context)

        # Extract table formatting
        table_fmt, cell_formats = self._extract_table_formatting(table)

        return TableData(
            headers=headers,
            rows=rows,
            table_type=table_type,
            context=context,
            cell_formats=cell_formats,
            table_format=table_fmt,
        )

    def _clean_cell_text(self, text: str) -> str:
        """Clean and normalize cell text."""
        # Remove extra whitespace and newlines
        cleaned = " ".join(text.split())
        return cleaned.strip()

    def _identify_table_type(
        self, headers: list[str], rows: list[list[str]], context: str
    ) -> str:
        """Identify the type of table based on headers and content."""
        headers_lower = [h.lower() for h in headers]
        headers_joined = " ".join(headers_lower)
        context_lower = context.lower()

        # Check for version history FIRST (very specific pattern)
        if "version" in headers_lower and ("date" in headers_joined or "revision" in headers_joined):
            return "version_history"

        # Check for terminology reference
        if "terminology" in headers_joined or "manch terminology" in headers_joined:
            return "terminology"

        # Check for integration mapping BEFORE field tables
        # Integration tables have headers like: Internal Field Name, External System, External Field Name
        if self._is_integration_table(headers_lower):
            return "integration_mapping"

        # Check for document requirements matrix
        if "documents" in headers_lower and len(headers) > 2:
            return "document_requirements"

        # Check for approval routing (but not if it's a field table with approver context)
        # Only classify as approval_routing if it doesn't look like a field table
        if any(kw in headers_joined for kw in ["tse", "asm"]) and not self._is_field_table(headers_lower, rows):
            return "approval_routing"

        # Check for field definition tables (now passes rows for validation)
        if self._is_field_table(headers_lower, rows):
            # Determine if this is the MASTER table or an actor-specific workflow table
            if self._is_master_field_table(context_lower):
                return "master_field_definitions"
            elif "initiator" in context_lower or "first party" in context_lower:
                return "initiator_fields"
            elif "vendor" in context_lower or "second party" in context_lower or "spoc" in context_lower:
                return "spoc_fields"
            elif "approver" in context_lower or "contributor" in context_lower:
                return "approver_fields"
            # Default: if no actor context and looks like field table, assume it's master
            return "master_field_definitions"

        # Check for dropdown/questionnaire tables
        if any("dropdown" in h.lower() or "question" in h.lower() for h in headers):
            return "dropdown_mapping"

        return "reference"

    def _is_integration_table(self, headers_lower: list[str]) -> bool:
        """
        Check if this is an integration mapping table.
        Integration tables have specific column patterns for system integration.
        """
        headers_joined = " ".join(headers_lower)

        # Check for integration-specific header patterns
        has_internal_field = any(
            kw in headers_joined
            for kw in ["internal field", "internal_field", "field name"]
        )
        has_external_system = any(
            kw in headers_joined
            for kw in ["external system", "external_system", "system"]
        )
        has_external_field = any(
            kw in headers_joined
            for kw in ["external field", "external_field", "ext_field"]
        )
        has_transformation = any(
            kw in headers_joined
            for kw in ["transformation", "mapping", "logic"]
        )

        # Integration table if it has internal + external field columns
        if has_internal_field and has_external_field:
            return True

        # Or if it has external system and field name
        if has_external_system and (has_external_field or has_transformation):
            return True

        return False

    def _is_master_field_table(self, context_lower: str) -> bool:
        """
        Determine if the table context indicates this is the MASTER field table
        (Section 4.4 Field-Level Information) vs a workflow/actor table.
        """
        # Check if context matches master table patterns
        for pattern in self.MASTER_TABLE_CONTEXT_PATTERNS:
            if pattern in context_lower:
                return True

        # Check if context matches workflow patterns (NOT master)
        for pattern in self.WORKFLOW_TABLE_CONTEXT_PATTERNS:
            if pattern in context_lower:
                return False

        # Default: if no clear context, treat as master (conservative approach)
        return True

    def _is_field_table(self, headers_lower: list[str], rows: list[list[str]] = None) -> bool:
        """
        Check if headers indicate a field definition table.
        Now more flexible - checks for any column with field names and valid field types.
        """
        # First try the strict patterns (for backward compatibility)
        for pattern in self.FIELD_HEADER_PATTERNS:
            if all(any(p in h for h in headers_lower) for p in pattern):
                return True

        # NEW: Flexible detection
        # Look for any column that might be field name (broader search)
        name_keywords = ["field", "name", "label", "attribute", "column"]
        type_keywords = ["type", "datatype", "data type", "field type"]

        has_name_column = any(
            any(keyword in h for keyword in name_keywords)
            for h in headers_lower
        )

        has_type_column = any(
            any(keyword in h for keyword in type_keywords)
            for h in headers_lower
        )

        # If we have both name and type columns, likely a field table
        if has_name_column and has_type_column:
            # Validate by checking if the type column contains valid field types
            if rows and self._validate_field_types_in_rows(headers_lower, rows):
                return True

        return False

    def _validate_field_types_in_rows(self, headers_lower: list[str], rows: list[list[str]]) -> bool:
        """
        Validate that rows contain valid field types.
        Checks if the TYPE column has values like TEXT, MOBILE, EMAIL, DATE, etc.
        """
        # Find type column
        type_col_idx = -1
        type_keywords = ["type", "datatype", "data type", "field type"]

        for i, header in enumerate(headers_lower):
            if any(keyword in header for keyword in type_keywords):
                type_col_idx = i
                break

        if type_col_idx == -1:
            return False

        # Valid field types (from FieldType enum + common variations)
        valid_types = {
            'text', 'dropdown', 'multi_dropdown', 'multidropdown', 'date', 'file',
            'checkbox', 'mobile', 'email', 'number', 'numeric', 'panel', 'label',
            'static_checkbox', 'external_dropdown', 'static check box',
            'external drop down value', 'multi dropdown',
            # Common document variations
            'string', 'int', 'integer', 'boolean', 'textarea', 'select'
        }

        # Check first 10 rows (or all if less than 10)
        valid_count = 0
        total_checked = 0

        for row in rows[:10]:
            if type_col_idx < len(row):
                cell_value = row[type_col_idx].lower().strip()
                if cell_value:  # Non-empty
                    total_checked += 1
                    # Check if it matches any valid type
                    if any(valid_type in cell_value for valid_type in valid_types):
                        valid_count += 1

        # If at least 50% of checked rows have valid field types, it's a field table
        if total_checked > 0 and (valid_count / total_checked) >= 0.5:
            return True

        return False

    def _extract_version_history(self, doc: Document, result: ParsedDocument):
        """Extract version history from document."""
        for table_data in result.raw_tables:
            if table_data.table_type == "version_history":
                headers_lower = [h.lower() for h in table_data.headers]

                for row in table_data.rows:
                    if not any(row):  # Skip empty rows
                        continue

                    entry = VersionEntry(
                        version=self._get_cell_by_header(row, headers_lower, ["version"]),
                        approved_by=self._get_cell_by_header(row, headers_lower, ["approved", "approver"]),
                        revision_date=self._get_cell_by_header(row, headers_lower, ["date", "revision"]),
                        description=self._get_cell_by_header(row, headers_lower, ["description", "change"]),
                        author=self._get_cell_by_header(row, headers_lower, ["author", "by"]),
                    )
                    if entry.version:  # Only add if version is present
                        result.version_history.append(entry)

    def _get_cell_by_header(
        self, row: list[str], headers_lower: list[str], search_terms: list[str]
    ) -> str:
        """Get cell value by matching header with search terms."""
        for term in search_terms:
            for i, header in enumerate(headers_lower):
                if term in header and i < len(row):
                    return row[i]
        return ""

    def _extract_terminology(self, doc: Document, result: ParsedDocument):
        """Extract terminology mappings."""
        for table_data in result.raw_tables:
            if table_data.table_type == "terminology":
                for row in table_data.rows:
                    if len(row) >= 2 and row[0] and row[1]:
                        result.terminology[row[0]] = row[1]

    def _extract_fields_from_tables(self, doc: Document, result: ParsedDocument):
        """
        Extract field definitions from tables with proper master/workflow distinction.

        Architecture:
        - MASTER table (Section 4.4 "Field-Level Information"): Contains ALL fields
          with complete logic. Only this table populates `all_fields`.
        - WORKFLOW tables (Section 4.5.x "Actor Behaviour"): Show which fields each
          actor sees with actor-specific overrides. Only populate actor-specific lists.

        Backward Compatibility:
        - If no master table exists (older BUD formats), actor tables will be used
          to populate `all_fields` as well, to maintain field visibility.

        This explicitly excludes:
        - Integration mapping tables
        - Approval routing tables
        - Document requirement matrices
        - Reference tables
        """
        current_panel = ""

        # Define valid field table types
        # MASTER table: the single source of truth for all field definitions
        MASTER_TABLE_TYPE = "master_field_definitions"

        # WORKFLOW tables: actor-specific views (subsets of master table)
        WORKFLOW_TABLE_TYPES = [
            "initiator_fields",
            "spoc_fields",
            "approver_fields",
        ]

        # Legacy support for old "field_definitions" type
        LEGACY_FIELD_TYPE = "field_definitions"

        ALL_VALID_TYPES = [MASTER_TABLE_TYPE, LEGACY_FIELD_TYPE] + WORKFLOW_TABLE_TYPES

        # Check if we have a master table
        has_master_table = any(
            t.table_type in [MASTER_TABLE_TYPE, LEGACY_FIELD_TYPE]
            for t in result.raw_tables
        )

        for table_data in result.raw_tables:
            # ONLY extract from explicitly defined field table types
            if table_data.table_type not in ALL_VALID_TYPES:
                continue

            headers_lower = [h.lower() for h in table_data.headers]
            table_type = table_data.table_type

            # Track previous field to detect consecutive duplicates
            previous_field = None
            previous_field_name = None

            for row in table_data.rows:
                field = self._parse_field_row(row, headers_lower, current_panel)

                if field:
                    # Track current panel for nested fields
                    if field.field_type == FieldType.PANEL:
                        current_panel = field.name

                    field_name_key = field.name.lower().strip()

                    # Check if this is a consecutive duplicate (same field name as previous row)
                    if previous_field and previous_field_name == field_name_key:
                        # This is a consecutive duplicate - merge the rules/logic
                        self._merge_field_data(previous_field, field)
                        # Don't add a new field, just merged into previous
                    else:
                        # New field - add it to the appropriate lists

                        # Determine if we should add to all_fields
                        should_add_to_all_fields = False
                        if table_type in [MASTER_TABLE_TYPE, LEGACY_FIELD_TYPE]:
                            # Master table always adds to all_fields
                            should_add_to_all_fields = True
                        elif not has_master_table:
                            # No master table: actor tables become the source
                            # Avoid duplicates by checking if field already exists
                            if not any(f.name.lower().strip() == field_name_key for f in result.all_fields):
                                should_add_to_all_fields = True

                        if should_add_to_all_fields:
                            result.all_fields.append(field)

                        # Add to actor-specific lists based on table type
                        if table_type == "initiator_fields":
                            result.initiator_fields.append(field)
                        elif table_type == "spoc_fields":
                            result.spoc_fields.append(field)
                        elif table_type == "approver_fields":
                            result.approver_fields.append(field)

                        # Update previous field tracking
                        previous_field = field
                        previous_field_name = field_name_key

    def _merge_field_data(self, existing_field: FieldDefinition, new_field: FieldDefinition):
        """Merge data from new_field into existing_field."""
        # Combine logic (only if new logic is different and not empty)
        if new_field.logic and new_field.logic.strip() != existing_field.logic.strip():
            if existing_field.logic:
                # Append new logic as a new bullet point
                existing_field.logic += "\n• " + new_field.logic
            else:
                existing_field.logic = new_field.logic

        # Combine rules (only if new rules are different and not empty)
        if new_field.rules and new_field.rules.strip() != existing_field.rules.strip():
            if existing_field.rules:
                existing_field.rules += "\n• " + new_field.rules
            else:
                existing_field.rules = new_field.rules

        # Merge dropdown values
        if new_field.dropdown_values:
            for val in new_field.dropdown_values:
                if val not in existing_field.dropdown_values:
                    existing_field.dropdown_values.append(val)

        # Update mandatory flag if new field is mandatory
        if new_field.is_mandatory:
            existing_field.is_mandatory = True

        # Merge visibility conditions
        if new_field.visibility_condition and new_field.visibility_condition.strip() != existing_field.visibility_condition.strip():
            if existing_field.visibility_condition:
                existing_field.visibility_condition += " | " + new_field.visibility_condition
            else:
                existing_field.visibility_condition = new_field.visibility_condition

        # Merge default values
        if new_field.default_value and not existing_field.default_value:
            existing_field.default_value = new_field.default_value

        # Merge validation
        if new_field.validation and new_field.validation.strip() != existing_field.validation.strip():
            if existing_field.validation:
                existing_field.validation += " | " + new_field.validation
            else:
                existing_field.validation = new_field.validation

    def _parse_field_row(
        self, row: list[str], headers_lower: list[str], current_panel: str
    ) -> Optional[FieldDefinition]:
        """Parse a single row into a FieldDefinition."""
        if not row or not any(row):
            return None

        # Find relevant columns with broader search terms
        # Name column - look for any column that might contain field names
        name_idx = self._find_column_index_flexible(
            headers_lower,
            ["field name", "filed name", "name", "field", "label", "attribute", "column name"]
        )

        # Type column - look for any column that might contain field types
        type_idx = self._find_column_index_flexible(
            headers_lower,
            ["field type", "type", "data type", "datatype", "field-type"]
        )

        # Mandatory column
        mandatory_idx = self._find_column_index_flexible(
            headers_lower,
            ["mandatory", "required", "is mandatory", "is required"]
        )

        # Logic/Rules column
        logic_idx = self._find_column_index_flexible(
            headers_lower,
            ["logic", "rules", "rule", "validation", "description", "notes", "logic and rules"]
        )

        if name_idx == -1 or type_idx == -1:
            return None

        name = row[name_idx] if name_idx < len(row) else ""
        field_type_raw = row[type_idx] if type_idx < len(row) else ""

        if not name:
            return None

        # Skip if name is just whitespace or empty
        if not name.strip():
            return None

        # Parse field type and skip if unknown
        field_type = FieldType.from_string(field_type_raw)
        if field_type == FieldType.UNKNOWN:
            return None

        mandatory_value = row[mandatory_idx] if mandatory_idx != -1 and mandatory_idx < len(row) else ""
        logic_value = row[logic_idx] if logic_idx != -1 and logic_idx < len(row) else ""

        # Parse mandatory flag
        is_mandatory = self._parse_mandatory(mandatory_value)

        # Extract visibility conditions and dropdown values from logic
        visibility_condition = ""
        dropdown_values = []

        if logic_value:
            visibility_match = re.search(
                r"(?:visible|show|display)\s+(?:if|when)\s+(.+?)(?:\.|$)",
                logic_value,
                re.IGNORECASE,
            )
            if visibility_match:
                visibility_condition = visibility_match.group(1).strip()

            # Extract dropdown values
            dropdown_match = re.search(
                r"(?:dropdown|values?)\s+(?:are|:)?\s*(.+?)(?:\.|$)",
                logic_value,
                re.IGNORECASE,
            )
            if dropdown_match:
                values_str = dropdown_match.group(1)
                # Split by common delimiters
                dropdown_values = [
                    v.strip()
                    for v in re.split(r"[,/]|(?:\s+and\s+)", values_str)
                    if v.strip()
                ]

        return FieldDefinition(
            name=name,
            field_type=field_type,
            field_type_raw=field_type_raw,
            is_mandatory=is_mandatory,
            logic=logic_value,
            section=current_panel,
            visibility_condition=visibility_condition,
            dropdown_values=dropdown_values,
        )

    def _find_column_index(self, headers: list[str], search_terms: list[str]) -> int:
        """Find column index matching any of the search terms."""
        for term in search_terms:
            for i, header in enumerate(headers):
                if term in header:
                    return i
        return -1

    def _find_column_index_flexible(self, headers: list[str], search_terms: list[str]) -> int:
        """
        Find column index with flexible matching.
        Tries multiple strategies to find the right column.
        """
        # Strategy 1: Exact match (prioritize exact matches)
        for term in search_terms:
            for i, header in enumerate(headers):
                if header.strip() == term:
                    return i

        # Strategy 2: Contains match (original behavior)
        for term in search_terms:
            for i, header in enumerate(headers):
                if term in header:
                    return i

        # Strategy 3: Partial word match (e.g., "type" matches "Type" or "Field-Type")
        for term in search_terms:
            term_words = set(term.split())
            for i, header in enumerate(headers):
                header_words = set(header.replace('-', ' ').replace('_', ' ').split())
                if term_words & header_words:  # If any words match
                    return i

        return -1

    def _parse_mandatory(self, value: str) -> bool:
        """Parse mandatory field value."""
        if not value:
            return False

        value_lower = value.lower().strip()

        if value_lower in ["yes", "true", "mandatory", "y", "1"]:
            return True
        if value_lower in ["no", "false", "optional", "n", "0"]:
            return False

        # Check for specific mandatory mentions
        if "mandatory" in value_lower:
            return True

        return False

    def _extract_workflows(self, doc: Document, result: ParsedDocument):
        """Extract workflow steps from document sections."""
        current_actor = ""
        step_number = 0

        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""
            text = para.text.strip()

            if not text:
                continue

            # Identify actor sections
            if "Heading" in style_name:
                actor = self._identify_actor(text)
                if actor:
                    current_actor = actor
                    step_number = 0
                    if actor not in result.workflows:
                        result.workflows[actor] = []

            # Parse workflow steps (list paragraphs)
            elif "List" in style_name and current_actor:
                step_number += 1
                step = WorkflowStep(
                    step_number=step_number,
                    description=text,
                    actor=current_actor,
                    action_type=self._identify_action_type(text),
                )

                # Check for notes
                if text.lower().startswith("note"):
                    step.notes.append(text)

                result.workflows[current_actor].append(step)

    def _identify_actor(self, text: str) -> str:
        """Identify actor from section heading."""
        text_lower = text.lower()

        if any(kw in text_lower for kw in self.INITIATOR_KEYWORDS):
            return "initiator"
        if any(kw in text_lower for kw in self.SPOC_KEYWORDS):
            return "spoc"
        if any(kw in text_lower for kw in self.APPROVAL_KEYWORDS):
            return "approver"

        return ""

    def _identify_action_type(self, text: str) -> str:
        """Identify action type from step description."""
        text_lower = text.lower()

        action_patterns = [
            ("login", ["login", "log in", "logs in"]),
            ("upload", ["upload", "uploads"]),
            ("validate", ["validate", "validates", "validation"]),
            ("approve", ["approve", "approves", "approval"]),
            ("reject", ["reject", "rejects", "rejection"]),
            ("submit", ["submit", "submits", "submission"]),
            ("notify", ["sms", "notification", "notify", "email"]),
            ("verify", ["verify", "verifies", "verification", "otp"]),
            ("create", ["create", "creates", "creation"]),
            ("update", ["update", "updates", "modify"]),
        ]

        for action_type, keywords in action_patterns:
            if any(kw in text_lower for kw in keywords):
                return action_type

        return ""

    def _extract_approval_rules(self, doc: Document, result: ParsedDocument):
        """Extract approval routing rules."""
        for table_data in result.raw_tables:
            if table_data.table_type == "approval_routing":
                for row in table_data.rows:
                    if len(row) >= 2:
                        rule = ApprovalRule(
                            condition=" | ".join(row[:-1]),
                            approver=row[-1] if row else "",
                        )
                        result.approval_rules.append(rule)

    def _extract_scope(self, doc: Document, result: ParsedDocument):
        """Extract scope, objectives, assumptions, and dependencies."""
        current_section = ""

        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""
            text = para.text.strip()

            if not text:
                continue

            # Track section context
            if "Heading" in style_name:
                text_lower = text.lower()
                if "in scope" in text_lower or "in-scope" in text_lower:
                    current_section = "scope_in"
                elif "out" in text_lower and "scope" in text_lower:
                    current_section = "scope_out"
                elif "objective" in text_lower:
                    current_section = "objectives"
                elif "assumption" in text_lower:
                    current_section = "assumptions"
                elif "dependenc" in text_lower:
                    current_section = "dependencies"
                else:
                    current_section = ""

            elif current_section and "List" in style_name:
                target_list = getattr(result, current_section, None)
                if target_list is not None:
                    target_list.append(text)

    def _extract_reference_tables(self, doc: Document, result: ParsedDocument):
        """Extract reference/lookup tables from Word document."""
        for table_data in result.raw_tables:
            if table_data.table_type == "reference":
                result.reference_tables.append(table_data)

        # Also extract tables from embedded Excel files
        self._extract_excel_reference_tables(result.file_path, result)

    def _extract_integration_info(self, doc: Document, result: ParsedDocument):
        """Extract integration field mappings."""
        for table_data in result.raw_tables:
            if table_data.table_type == "integration_mapping":
                headers_lower = [h.lower() for h in table_data.headers]

                for row in table_data.rows:
                    if not any(row):
                        continue

                    field = IntegrationField(
                        internal_field_name=self._get_cell_by_header(
                            row, headers_lower, ["internal", "field name"]
                        ),
                        external_system=self._get_cell_by_header(
                            row, headers_lower, ["external system", "system"]
                        ),
                        external_field_name=self._get_cell_by_header(
                            row, headers_lower, ["external field"]
                        ),
                        data_type=self._get_cell_by_header(
                            row, headers_lower, ["data type", "type"]
                        ),
                        transformation_logic=self._get_cell_by_header(
                            row, headers_lower, ["transformation", "logic"]
                        ),
                        is_mandatory=self._parse_mandatory(
                            self._get_cell_by_header(row, headers_lower, ["mandatory"])
                        ),
                        default_value=self._get_cell_by_header(
                            row, headers_lower, ["default"]
                        ),
                        validation_rules=self._get_cell_by_header(
                            row, headers_lower, ["validation"]
                        ),
                    )
                    result.integration_fields.append(field)

    def _extract_document_requirements(self, doc: Document, result: ParsedDocument):
        """Extract document requirements matrix."""
        for table_data in result.raw_tables:
            if table_data.table_type == "document_requirements":
                if len(table_data.headers) < 2:
                    continue

                # First column is document name, rest are vendor types
                vendor_types = table_data.headers[1:]

                # Handle merged header cells
                if len(table_data.rows) > 0:
                    # Check if first row contains vendor type codes
                    first_row = table_data.rows[0]
                    if first_row and all(
                        any(c.isupper() for c in cell) for cell in first_row[1:] if cell
                    ):
                        vendor_types = first_row[1:]
                        table_data.rows = table_data.rows[1:]

                for row in table_data.rows:
                    if not row or not row[0]:
                        continue

                    doc_name = row[0]
                    requirements = {}

                    for i, vendor_type in enumerate(vendor_types):
                        if i + 1 < len(row):
                            requirements[vendor_type] = row[i + 1]

                    result.document_requirements.append(
                        DocumentRequirementMatrix(
                            document_name=doc_name,
                            requirements=requirements,
                        )
                    )

    def _extract_dropdown_values(self, result: ParsedDocument):
        """Consolidate dropdown values from all fields."""
        for field in result.all_fields:
            if field.dropdown_values:
                result.dropdown_mappings[field.name] = field.dropdown_values

        # Also extract from reference tables that look like dropdowns
        for table_data in result.raw_tables:
            if table_data.table_type == "dropdown_mapping":
                for i, header in enumerate(table_data.headers):
                    values = []
                    for row in table_data.rows:
                        if i < len(row) and row[i]:
                            values.append(row[i])
                    if values:
                        result.dropdown_mappings[header] = values

    def _extract_images(self, file_path: str, doc: Document) -> list[ImageReference]:
        """Extract all images from the document with their metadata."""
        images = []
        position_index = 0

        try:
            # Open DOCX as ZIP file
            with zipfile.ZipFile(file_path, 'r') as docx_zip:
                # List all files in word/media/
                media_files = [f for f in docx_zip.namelist() if f.startswith('word/media/')]

                for media_file in media_files:
                    # Read binary data
                    image_data = docx_zip.read(media_file)

                    # Encode as base64
                    image_data_base64 = base64.b64encode(image_data).decode('utf-8')

                    # Determine content type from extension
                    filename = media_file.split('/')[-1]
                    ext = filename.split('.')[-1].lower()
                    content_type_map = {
                        'png': 'image/png',
                        'jpg': 'image/jpeg',
                        'jpeg': 'image/jpeg',
                        'gif': 'image/gif',
                        'bmp': 'image/bmp',
                        'tiff': 'image/tiff',
                        'emf': 'image/x-emf',
                        'wmf': 'image/x-wmf',
                    }
                    content_type = content_type_map.get(ext, 'image/unknown')

                    # Try to get dimensions from inline shapes
                    width_inches = 1.0  # Default
                    height_inches = 1.0  # Default

                    # Check if this image is referenced in inline shapes
                    for shape in doc.inline_shapes:
                        try:
                            # Get the relationship ID
                            if hasattr(shape, '_inline') and hasattr(shape._inline, 'graphic'):
                                # Try to match the image
                                # Default dimensions from shape
                                width_inches = shape.width.inches if hasattr(shape.width, 'inches') else 1.0
                                height_inches = shape.height.inches if hasattr(shape.height, 'inches') else 1.0
                                break
                        except:
                            pass

                    images.append(ImageReference(
                        filename=filename,
                        image_data_base64=image_data_base64,
                        width_inches=width_inches,
                        height_inches=height_inches,
                        content_type=content_type,
                        position_index=position_index,
                    ))
                    position_index += 1

        except Exception as e:
            # Log error but don't fail the entire parsing
            print(f"Warning: Could not extract images: {e}")

        return images

    def _extract_run_formatting(self, paragraph) -> list[RunFormatting]:
        """Extract formatted text runs from a paragraph."""
        run_list = []

        try:
            for run in paragraph.runs:
                # Extract font information
                font = run.font
                font_info = FontInfo(
                    family=font.name or "",
                    size_pt=font.size.pt if font.size else 0.0,
                    bold=font.bold if font.bold is not None else False,
                    italic=font.italic if font.italic is not None else False,
                    underline=font.underline if font.underline is not None else False,
                    color_rgb=None,
                )

                # Extract color if available
                if font.color and font.color.rgb:
                    try:
                        rgb = font.color.rgb
                        font_info.color_rgb = (rgb[0], rgb[1], rgb[2])
                    except:
                        pass

                run_list.append(RunFormatting(
                    text=run.text,
                    font=font_info,
                ))
        except Exception as e:
            # If extraction fails, return empty list
            pass

        return run_list

    def _extract_paragraph_formatting(self, paragraph) -> ParagraphFormatting:
        """Extract paragraph formatting information."""
        try:
            fmt = paragraph.paragraph_format

            # Map alignment enum to string
            alignment_map = {
                WD_ALIGN_PARAGRAPH.LEFT: "left",
                WD_ALIGN_PARAGRAPH.CENTER: "center",
                WD_ALIGN_PARAGRAPH.RIGHT: "right",
                WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
                None: "left",
            }

            alignment = alignment_map.get(fmt.alignment, "left")

            return ParagraphFormatting(
                alignment=alignment,
                line_spacing=fmt.line_spacing if fmt.line_spacing else 1.0,
                space_before_pt=fmt.space_before.pt if fmt.space_before else 0.0,
                space_after_pt=fmt.space_after.pt if fmt.space_after else 0.0,
                left_indent_pt=fmt.left_indent.pt if fmt.left_indent else 0.0,
                right_indent_pt=fmt.right_indent.pt if fmt.right_indent else 0.0,
                first_line_indent_pt=fmt.first_line_indent.pt if fmt.first_line_indent else 0.0,
            )
        except Exception as e:
            # Return default formatting if extraction fails
            return ParagraphFormatting()

    def _extract_page_setup(self, doc: Document) -> PageSetup:
        """Extract page setup information."""
        try:
            section = doc.sections[0]  # Get first section's page setup

            return PageSetup(
                page_width_inches=section.page_width.inches if section.page_width else 8.5,
                page_height_inches=section.page_height.inches if section.page_height else 11.0,
                left_margin_inches=section.left_margin.inches if section.left_margin else 1.0,
                right_margin_inches=section.right_margin.inches if section.right_margin else 1.0,
                top_margin_inches=section.top_margin.inches if section.top_margin else 1.0,
                bottom_margin_inches=section.bottom_margin.inches if section.bottom_margin else 1.0,
                orientation="landscape" if section.orientation == 1 else "portrait"
            )
        except Exception as e:
            print(f"Warning: Could not extract page setup: {e}")
            return PageSetup()

    def _extract_headers_footers(self, doc: Document, file_path: str) -> tuple[Optional[HeaderFooter], Optional[HeaderFooter]]:
        """Extract header and footer content."""
        header_content = None
        footer_content = None

        try:
            section = doc.sections[0]

            # Extract header
            if section.header:
                header_content = HeaderFooter()
                for para in section.header.paragraphs:
                    text = para.text.strip()
                    header_content.paragraphs.append(text)
                    header_content.runs.append(self._extract_run_formatting(para))
                    header_content.paragraph_formats.append(self._extract_paragraph_formatting(para))

                # Extract images from header
                try:
                    with zipfile.ZipFile(file_path, 'r') as docx_zip:
                        # Check for images in header
                        media_files = [f for f in docx_zip.namelist() if f.startswith('word/media/')]
                        # We'll extract all images and they'll appear in the header if they're there
                except:
                    pass

            # Extract footer
            if section.footer:
                footer_content = HeaderFooter()
                for para in section.footer.paragraphs:
                    text = para.text.strip()
                    footer_content.paragraphs.append(text)
                    footer_content.runs.append(self._extract_run_formatting(para))
                    footer_content.paragraph_formats.append(self._extract_paragraph_formatting(para))

        except Exception as e:
            print(f"Warning: Could not extract headers/footers: {e}")

        return header_content, footer_content

    def _extract_exact_document_order(self, doc: Document, file_path: str) -> list[DocumentElement]:
        """
        Extract the EXACT order of all document elements.
        This preserves the exact structure for perfect recreation.
        """
        elements = []
        index = 0
        image_index = 0

        # Build a mapping of images by filename
        images_map = {}
        try:
            with zipfile.ZipFile(file_path, 'r') as docx_zip:
                media_files = [f for f in docx_zip.namelist() if f.startswith('word/media/')]
                for media_file in media_files:
                    filename = media_file.split('/')[-1]
                    image_data = docx_zip.read(media_file)
                    image_data_base64 = base64.b64encode(image_data).decode('utf-8')
                    images_map[filename] = image_data_base64
        except:
            pass

        # Iterate through document body in exact order
        for element in doc.element.body:
            if element.tag.endswith("p"):
                # Find corresponding paragraph
                para = None
                for p in doc.paragraphs:
                    if p._element is element:
                        para = p
                        break

                if para is None:
                    continue

                text = para.text  # Don't strip - preserve exact spacing
                style_name = para.style.name if para.style else ""
                heading_level = self._get_heading_level(style_name)

                # Check if paragraph contains images
                has_images = False
                for run in para.runs:
                    if hasattr(run, '_element'):
                        for drawing in run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                            has_images = True
                            # Extract image reference
                            try:
                                for inline in doc.inline_shapes:
                                    width_inches = inline.width.inches if hasattr(inline.width, 'inches') else 1.0
                                    height_inches = inline.height.inches if hasattr(inline.height, 'inches') else 1.0

                                    # Create image element
                                    img_elem = DocumentElement(
                                        element_type="image",
                                        index=index,
                                        content=ImageReference(
                                            filename=f"image{image_index+1}",
                                            image_data_base64="",  # Will be filled later
                                            width_inches=width_inches,
                                            height_inches=height_inches,
                                            content_type="image/unknown",
                                            position_index=image_index
                                        )
                                    )
                                    elements.append(img_elem)
                                    index += 1
                                    image_index += 1
                                    break
                            except:
                                pass

                # DON'T skip empty paragraphs - they're important for spacing!
                # Extract formatting
                runs = self._extract_run_formatting(para)
                para_format = self._extract_paragraph_formatting(para)

                if heading_level > 0:
                    # This is a heading
                    elem = DocumentElement(
                        element_type="heading",
                        index=index,
                        content=text,
                        runs=runs,
                        paragraph_format=para_format,
                        heading_level=heading_level,
                        style_name=style_name
                    )
                else:
                    # Regular paragraph
                    elem = DocumentElement(
                        element_type="paragraph",
                        index=index,
                        content=text,
                        runs=runs,
                        paragraph_format=para_format,
                        style_name=style_name
                    )

                elements.append(elem)
                index += 1

            elif element.tag.endswith("tbl"):
                # Find corresponding table
                table = None
                for tbl in doc.tables:
                    if tbl._tbl is element:
                        table = tbl
                        break

                if table is None:
                    continue

                # Parse table with formatting
                table_data = self._parse_table(table, "")

                elem = DocumentElement(
                    element_type="table",
                    index=index,
                    content=table_data
                )

                elements.append(elem)
                index += 1

        return elements

    def _extract_table_formatting(self, table: Table) -> tuple[TableFormatting, list[list[CellFormatting]]]:
        """Extract table and cell formatting."""
        try:
            # Table-level formatting
            table_fmt = TableFormatting(
                alignment="left",  # Default, python-docx doesn't expose table alignment easily
                width_inches=None,
            )

            # Cell formatting matrix
            cell_formats = []
            for row in table.rows:
                row_formats = []
                for cell in row.cells:
                    cell_fmt = CellFormatting(
                        background_color=None,
                        vertical_alignment="top",
                    )

                    # Try to extract cell shading/background
                    try:
                        if hasattr(cell, '_element') and hasattr(cell._element, 'tcPr'):
                            tc_pr = cell._element.tcPr
                            if tc_pr is not None:
                                shd = tc_pr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                                if shd is not None and 'fill' in shd.attrib:
                                    fill_color = shd.attrib['fill']
                                    # Convert hex color to RGB tuple
                                    if fill_color and fill_color != 'auto' and len(fill_color) == 6:
                                        try:
                                            r = int(fill_color[0:2], 16)
                                            g = int(fill_color[2:4], 16)
                                            b = int(fill_color[4:6], 16)
                                            cell_fmt.background_color = (r, g, b)
                                        except:
                                            pass
                    except:
                        pass

                    row_formats.append(cell_fmt)
                cell_formats.append(row_formats)

            return table_fmt, cell_formats
        except Exception as e:
            # Return defaults if extraction fails
            return TableFormatting(), []

    def _extract_excel_reference_tables(self, file_path: str, result: ParsedDocument):
        """
        Extract reference tables from embedded Excel files in the DOCX.

        Args:
            file_path: Path to the DOCX file
            result: ParsedDocument to add Excel tables to
        """
        try:
            import openpyxl
            from openpyxl.utils.exceptions import InvalidFileException
            import tempfile
            import os

            print("\nSearching for embedded Excel files...")

            with zipfile.ZipFile(file_path, 'r') as docx_zip:
                # List all files in the DOCX archive
                all_files = docx_zip.namelist()

                # Look for Excel files in embeddings directory
                excel_files = [
                    f for f in all_files
                    if f.startswith('word/embeddings/') and (f.endswith('.xlsx') or f.endswith('.xls'))
                ]

                if not excel_files:
                    print("  No embedded Excel files found")
                    return

                print(f"  Found {len(excel_files)} embedded Excel file(s)")

                for excel_file in excel_files:
                    try:
                        # Extract filename
                        filename = excel_file.split('/')[-1]
                        print(f"\n  Processing: {filename}")

                        # Extract Excel file to temporary location
                        excel_data = docx_zip.read(excel_file)

                        # Create temporary file
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as temp_file:
                            temp_file.write(excel_data)
                            temp_path = temp_file.name

                        try:
                            # Open Excel file with openpyxl
                            wb = openpyxl.load_workbook(temp_path, data_only=True)

                            # Process each sheet
                            for sheet_name in wb.sheetnames:
                                sheet = wb[sheet_name]
                                print(f"    Sheet: {sheet_name}")

                                # Extract table data from sheet
                                table_data = self._parse_excel_sheet(sheet, filename, sheet_name)

                                if table_data and (table_data.headers or table_data.rows):
                                    # Mark as reference table from Excel
                                    table_data.table_type = "reference"
                                    table_data.source = "excel"
                                    table_data.source_file = filename
                                    table_data.sheet_name = sheet_name

                                    # Add to reference tables
                                    result.reference_tables.append(table_data)
                                    print(f"      ✓ Extracted table: {len(table_data.headers)} columns, {len(table_data.rows)} rows")
                                else:
                                    print(f"      ○ Sheet is empty or has no data")

                            wb.close()

                        finally:
                            # Clean up temp file
                            try:
                                os.unlink(temp_path)
                            except:
                                pass

                    except InvalidFileException as e:
                        print(f"    ✗ Could not open Excel file: {e}")
                    except Exception as e:
                        print(f"    ✗ Error processing {filename}: {e}")

        except ImportError:
            print("  ⚠ openpyxl not installed - cannot extract Excel files")
            print("  Install with: pip install openpyxl")
        except Exception as e:
            print(f"  ✗ Error extracting Excel files: {e}")

    def _parse_excel_sheet(self, sheet, filename: str, sheet_name: str) -> Optional[TableData]:
        """
        Parse an Excel sheet into TableData.

        Args:
            sheet: openpyxl worksheet
            filename: Excel filename
            sheet_name: Sheet name

        Returns:
            TableData or None if sheet is empty
        """
        try:
            # Get all rows as values
            data = list(sheet.values)

            if not data:
                return None

            # Filter out completely empty rows
            data = [row for row in data if any(cell is not None and str(cell).strip() for cell in row)]

            if not data:
                return None

            # First row is headers
            headers = []
            if data:
                headers = [str(cell) if cell is not None else "" for cell in data[0]]
                data = data[1:]

            # Convert remaining rows
            rows = []
            for row in data:
                # Convert all cells to strings, handling None
                row_data = [str(cell) if cell is not None else "" for cell in row]
                # Pad or trim to match header length
                while len(row_data) < len(headers):
                    row_data.append("")
                row_data = row_data[:len(headers)]
                rows.append(row_data)

            return TableData(
                headers=headers,
                rows=rows,
                table_type="reference",
                context=f"Excel: {filename} - {sheet_name}",
                source="excel",
                source_file=filename,
                sheet_name=sheet_name,
            )

        except Exception as e:
            print(f"    ✗ Error parsing sheet: {e}")
            return None
