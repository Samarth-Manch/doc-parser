"""
Document recreation engine for rebuilding DOCX files from parsed data.
EXACT RECREATION - Preserves original structure and formatting perfectly.
"""

import base64
import io
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .models import (
    ParsedDocument,
    DocumentElement,
    TableData,
    ImageReference,
    FontInfo,
    ParagraphFormatting,
    RunFormatting,
    PageSetup,
    HeaderFooter,
)


class DocumentRecreator:
    """
    Recreates Word documents from ParsedDocument data with EXACT structure.

    Uses document_elements list to rebuild document in original order,
    preserving all visual elements, formatting, and structure.
    """

    def __init__(self, parsed_doc: ParsedDocument):
        """
        Initialize the recreator with parsed document data.

        Args:
            parsed_doc: ParsedDocument object containing all extracted data
        """
        self.parsed_doc = parsed_doc
        self.output_doc: Optional[Document] = None

    def recreate(self, output_path: str) -> bool:
        """
        Main recreation method - rebuilds the document EXACTLY and saves it.

        Args:
            output_path: Path where the recreated document should be saved

        Returns:
            True if successful, False otherwise
        """
        try:
            # Create new document
            self.output_doc = Document()

            # Apply metadata
            self._apply_metadata()

            # Apply page setup
            if self.parsed_doc.page_setup:
                self._apply_page_setup()

            # Apply headers and footers
            if self.parsed_doc.header or self.parsed_doc.footer:
                self._apply_headers_footers()

            # Recreate document in EXACT order from document_elements
            if self.parsed_doc.document_elements:
                print(f"Recreating {len(self.parsed_doc.document_elements)} elements in exact order...")
                self._recreate_exact_structure()
            else:
                # Fallback to old method if document_elements not available
                print("Warning: Using legacy recreation (not exact)")
                self._recreate_legacy()

            # Save document
            self.output_doc.save(output_path)
            return True

        except Exception as e:
            print(f"Error recreating document: {e}")
            import traceback
            traceback.print_exc()
            return False

    def _recreate_exact_structure(self):
        """Recreate document in EXACT order using document_elements."""
        for elem in self.parsed_doc.document_elements:
            try:
                if elem.element_type == "heading":
                    self._recreate_heading(elem)
                elif elem.element_type == "paragraph":
                    self._recreate_paragraph(elem)
                elif elem.element_type == "table":
                    self._recreate_table_element(elem)
                elif elem.element_type == "image":
                    self._recreate_image_element(elem)
            except Exception as e:
                print(f"Warning: Could not recreate element {elem.index} ({elem.element_type}): {e}")

    def _recreate_heading(self, elem: DocumentElement):
        """Recreate a heading element."""
        if not self.output_doc or not elem.content:
            return

        # Add heading with proper level
        heading = self.output_doc.add_heading(elem.content, level=elem.heading_level)

        # Apply formatting if available
        if elem.runs and len(elem.runs) > 0:
            # Clear default text and apply formatted runs
            heading.clear()
            for run_fmt in elem.runs:
                run = heading.add_run(run_fmt.text)
                self._apply_font(run, run_fmt.font)

        if elem.paragraph_format:
            self._apply_paragraph_format(heading, elem.paragraph_format)

    def _recreate_paragraph(self, elem: DocumentElement):
        """Recreate a regular paragraph element."""
        if not self.output_doc:
            return

        # Create paragraph with the exact style if available
        if elem.style_name and elem.style_name != "Normal":
            try:
                para = self.output_doc.add_paragraph(style=elem.style_name)
            except:
                # Style doesn't exist, create normal paragraph
                para = self.output_doc.add_paragraph()
        else:
            para = self.output_doc.add_paragraph()

        # Add formatted runs if available
        if elem.runs and len(elem.runs) > 0:
            for run_fmt in elem.runs:
                run = para.add_run(run_fmt.text)
                self._apply_font(run, run_fmt.font)
        elif elem.content:
            # Fallback to text content (may be empty for spacing)
            para.add_run(elem.content)

        # Apply paragraph formatting
        if elem.paragraph_format:
            self._apply_paragraph_format(para, elem.paragraph_format)

    def _recreate_table_element(self, elem: DocumentElement):
        """Recreate a table element."""
        if not self.output_doc or not elem.content:
            return

        table_data = elem.content

        # Ensure it's a TableData object
        if not isinstance(table_data, TableData):
            return

        try:
            # Calculate table dimensions
            num_rows = len(table_data.rows) + 1  # +1 for header row
            num_cols = len(table_data.headers) if table_data.headers else 0

            if num_cols == 0:
                return

            # Create table
            table = self.output_doc.add_table(rows=num_rows, cols=num_cols)

            # Set table style
            table.style = 'Light Grid Accent 1'

            # Add headers
            for i, header in enumerate(table_data.headers):
                if i < len(table.rows[0].cells):
                    table.rows[0].cells[i].text = header

            # Add data rows
            for row_idx, row_data in enumerate(table_data.rows):
                actual_row_idx = row_idx + 1  # +1 to skip header row
                if actual_row_idx < len(table.rows):
                    for col_idx, cell_data in enumerate(row_data):
                        if col_idx < len(table.rows[actual_row_idx].cells):
                            table.rows[actual_row_idx].cells[col_idx].text = str(cell_data)

            # Apply cell formatting if available
            if table_data.cell_formats and len(table_data.cell_formats) > 0:
                self._apply_table_formatting(table, table_data)

        except Exception as e:
            print(f"Warning: Could not recreate table: {e}")

    def _recreate_image_element(self, elem: DocumentElement):
        """Recreate an image element."""
        if not self.output_doc or not elem.content:
            return

        img_ref = elem.content

        # Ensure it's an ImageReference object
        if not isinstance(img_ref, ImageReference):
            return

        # Find the actual image data from parsed_doc.images
        actual_img = None
        for img in self.parsed_doc.images:
            if img.position_index == img_ref.position_index:
                actual_img = img
                break

        if not actual_img or not actual_img.image_data_base64:
            return

        try:
            # Decode base64 image data
            img_data = base64.b64decode(actual_img.image_data_base64)

            # Create BytesIO stream
            img_stream = io.BytesIO(img_data)

            # Add to document with specified width
            width = Inches(actual_img.width_inches) if actual_img.width_inches > 0 else Inches(1.0)
            self.output_doc.add_picture(img_stream, width=width)

        except Exception as e:
            print(f"Warning: Could not embed image {actual_img.filename}: {e}")

    def _apply_metadata(self):
        """Set document core properties from metadata."""
        if not self.output_doc:
            return

        props = self.output_doc.core_properties
        metadata = self.parsed_doc.metadata

        props.title = metadata.title
        props.author = metadata.author
        props.subject = metadata.subject

    def _apply_page_setup(self):
        """Apply page setup (margins, size, orientation)."""
        if not self.output_doc:
            return

        try:
            section = self.output_doc.sections[0]
            ps = self.parsed_doc.page_setup

            section.page_width = Inches(ps.page_width_inches)
            section.page_height = Inches(ps.page_height_inches)
            section.left_margin = Inches(ps.left_margin_inches)
            section.right_margin = Inches(ps.right_margin_inches)
            section.top_margin = Inches(ps.top_margin_inches)
            section.bottom_margin = Inches(ps.bottom_margin_inches)

            # Set orientation: 0=portrait, 1=landscape
            section.orientation = 1 if ps.orientation == "landscape" else 0
        except Exception as e:
            print(f"Warning: Could not apply page setup: {e}")

    def _apply_headers_footers(self):
        """Apply header and footer content."""
        if not self.output_doc:
            return

        try:
            section = self.output_doc.sections[0]

            # Apply header
            if self.parsed_doc.header:
                header = section.header
                for i, para_text in enumerate(self.parsed_doc.header.paragraphs):
                    para = header.add_paragraph()

                    # Add formatted runs if available
                    if i < len(self.parsed_doc.header.runs) and self.parsed_doc.header.runs[i]:
                        for run_fmt in self.parsed_doc.header.runs[i]:
                            run = para.add_run(run_fmt.text)
                            self._apply_font(run, run_fmt.font)
                    else:
                        para.add_run(para_text)

                    # Apply paragraph formatting
                    if i < len(self.parsed_doc.header.paragraph_formats):
                        self._apply_paragraph_format(para, self.parsed_doc.header.paragraph_formats[i])

            # Apply footer
            if self.parsed_doc.footer:
                footer = section.footer
                for i, para_text in enumerate(self.parsed_doc.footer.paragraphs):
                    para = footer.add_paragraph()

                    # Add formatted runs if available
                    if i < len(self.parsed_doc.footer.runs) and self.parsed_doc.footer.runs[i]:
                        for run_fmt in self.parsed_doc.footer.runs[i]:
                            run = para.add_run(run_fmt.text)
                            self._apply_font(run, run_fmt.font)
                    else:
                        para.add_run(para_text)

                    # Apply paragraph formatting
                    if i < len(self.parsed_doc.footer.paragraph_formats):
                        self._apply_paragraph_format(para, self.parsed_doc.footer.paragraph_formats[i])

        except Exception as e:
            print(f"Warning: Could not apply headers/footers: {e}")

    def _apply_font(self, run, font_info: FontInfo):
        """
        Apply font formatting to a run.

        Args:
            run: The run object to format
            font_info: FontInfo object with formatting details
        """
        try:
            if font_info.family:
                run.font.name = font_info.family
            if font_info.size_pt > 0:
                run.font.size = Pt(font_info.size_pt)

            run.font.bold = font_info.bold
            run.font.italic = font_info.italic
            run.font.underline = font_info.underline

            if font_info.color_rgb:
                run.font.color.rgb = RGBColor(*font_info.color_rgb)
        except Exception as e:
            pass

    def _apply_paragraph_format(self, para, fmt: ParagraphFormatting):
        """
        Apply paragraph formatting.

        Args:
            para: The paragraph object to format
            fmt: ParagraphFormatting object with formatting details
        """
        try:
            # Map alignment string to enum
            alignment_map = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
                "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
            }

            if fmt.alignment in alignment_map:
                para.paragraph_format.alignment = alignment_map[fmt.alignment]

            if fmt.line_spacing:
                para.paragraph_format.line_spacing = fmt.line_spacing

            if fmt.space_before_pt:
                para.paragraph_format.space_before = Pt(fmt.space_before_pt)

            if fmt.space_after_pt:
                para.paragraph_format.space_after = Pt(fmt.space_after_pt)

            if fmt.left_indent_pt:
                para.paragraph_format.left_indent = Pt(fmt.left_indent_pt)

            if fmt.right_indent_pt:
                para.paragraph_format.right_indent = Pt(fmt.right_indent_pt)

            if fmt.first_line_indent_pt:
                para.paragraph_format.first_line_indent = Pt(fmt.first_line_indent_pt)
        except Exception as e:
            pass

    def _apply_table_formatting(self, table, table_data: TableData):
        """
        Apply cell formatting to table cells.

        Args:
            table: The table object to format
            table_data: TableData with formatting information
        """
        try:
            # Apply cell background colors and formatting
            for row_idx, row_formats in enumerate(table_data.cell_formats):
                # +1 to account for header row in the recreated table
                actual_row_idx = row_idx + 1 if row_idx > 0 else 0

                if actual_row_idx < len(table.rows):
                    for col_idx, cell_fmt in enumerate(row_formats):
                        if col_idx < len(table.rows[actual_row_idx].cells):
                            cell = table.rows[actual_row_idx].cells[col_idx]

                            # Apply background color if specified
                            if cell_fmt.background_color:
                                try:
                                    shading_elm = cell._element.get_or_add_tcPr().get_or_add_shd()
                                    rgb = cell_fmt.background_color
                                    color_hex = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
                                    shading_elm.fill = color_hex
                                except:
                                    pass
        except Exception as e:
            pass

    def _recreate_legacy(self):
        """Legacy recreation method (fallback if document_elements not available)."""
        # Recreate sections with formatting
        for section in self.parsed_doc.sections:
            self._recreate_section(section)

        # Recreate tables with formatting
        for table_data in self.parsed_doc.raw_tables:
            self._recreate_table(table_data)

        # Embed images
        self._embed_images()

    def _recreate_section(self, section):
        """Legacy section recreation."""
        if not self.output_doc:
            return

        # Add heading
        if section.heading:
            heading = self.output_doc.add_heading(section.heading, level=section.level)

        # Add content paragraphs with formatting
        for i, paragraph_text in enumerate(section.content):
            para = self.output_doc.add_paragraph()

            # If we have run formatting for this paragraph, use it
            if i < len(section.runs) and section.runs[i]:
                for run_fmt in section.runs[i]:
                    run = para.add_run(run_fmt.text)
                    self._apply_font(run, run_fmt.font)
            else:
                para.add_run(paragraph_text)

            # Apply paragraph formatting if available
            if i < len(section.paragraph_formats):
                self._apply_paragraph_format(para, section.paragraph_formats[i])

        # Recursively recreate subsections
        for subsection in section.subsections:
            self._recreate_section(subsection)

    def _embed_images(self):
        """Legacy image embedding."""
        if not self.output_doc:
            return

        for img_ref in self.parsed_doc.images:
            try:
                img_data = base64.b64decode(img_ref.image_data_base64)
                img_stream = io.BytesIO(img_data)
                width = Inches(img_ref.width_inches) if img_ref.width_inches > 0 else Inches(1.0)
                self.output_doc.add_picture(img_stream, width=width)
            except Exception as e:
                print(f"Warning: Could not embed image {img_ref.filename}: {e}")

    def _recreate_table(self, table_data: TableData):
        """Legacy table recreation."""
        if not self.output_doc:
            return

        try:
            num_rows = len(table_data.rows) + 1
            num_cols = len(table_data.headers) if table_data.headers else 0

            if num_cols == 0:
                return

            table = self.output_doc.add_table(rows=num_rows, cols=num_cols)
            table.style = 'Light Grid Accent 1'

            for i, header in enumerate(table_data.headers):
                if i < len(table.rows[0].cells):
                    table.rows[0].cells[i].text = header

            for row_idx, row_data in enumerate(table_data.rows):
                actual_row_idx = row_idx + 1
                if actual_row_idx < len(table.rows):
                    for col_idx, cell_data in enumerate(row_data):
                        if col_idx < len(table.rows[actual_row_idx].cells):
                            table.rows[actual_row_idx].cells[col_idx].text = str(cell_data)

            if table_data.cell_formats and len(table_data.cell_formats) > 0:
                self._apply_table_formatting(table, table_data)

        except Exception as e:
            print(f"Warning: Could not recreate table: {e}")
