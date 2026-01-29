"""
Script to explore the structure of OOXML documents.
"""

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from pathlib import Path
import json

def explore_document(doc_path: str) -> dict:
    """Explore the structure of a Word document."""
    doc = Document(doc_path)

    result = {
        "path": str(doc_path),
        "sections": [],
        "tables_count": len(doc.tables),
        "paragraphs_count": len(doc.paragraphs),
        "core_properties": {},
        "structure": []
    }

    # Extract core properties
    props = doc.core_properties
    result["core_properties"] = {
        "author": props.author,
        "title": props.title,
        "subject": props.subject,
        "created": str(props.created) if props.created else None,
        "modified": str(props.modified) if props.modified else None,
        "last_modified_by": props.last_modified_by,
    }

    # Track document structure (paragraphs and tables in order)
    for element in doc.element.body:
        if element.tag.endswith('p'):  # Paragraph
            para_elem = element
            # Find corresponding paragraph object
            for para in doc.paragraphs:
                if para._element is para_elem:
                    text = para.text.strip()
                    style = para.style.name if para.style else None
                    if text:  # Only include non-empty paragraphs
                        result["structure"].append({
                            "type": "paragraph",
                            "style": style,
                            "text": text[:200] + "..." if len(text) > 200 else text
                        })
                    break
        elif element.tag.endswith('tbl'):  # Table
            for table in doc.tables:
                if table._tbl is element:
                    rows = len(table.rows)
                    cols = len(table.columns) if table.rows else 0
                    # Get first row as header hint
                    header_cells = []
                    if table.rows:
                        for cell in table.rows[0].cells:
                            header_cells.append(cell.text.strip()[:50])
                    result["structure"].append({
                        "type": "table",
                        "rows": rows,
                        "cols": cols,
                        "header_hint": header_cells
                    })
                    break

    return result


def main():
    docs_dir = Path("documents")

    for doc_path in sorted(docs_dir.glob("*.docx")):
        print(f"\n{'='*80}")
        print(f"Document: {doc_path.name}")
        print('='*80)

        try:
            result = explore_document(doc_path)

            print(f"\nMetadata:")
            print(f"  - Paragraphs: {result['paragraphs_count']}")
            print(f"  - Tables: {result['tables_count']}")

            print(f"\nCore Properties:")
            for key, value in result['core_properties'].items():
                if value:
                    print(f"  - {key}: {value}")

            print(f"\nDocument Structure:")
            for i, item in enumerate(result['structure'][:30]):  # First 30 elements
                if item['type'] == 'paragraph':
                    print(f"  [{i}] PARA ({item['style']}): {item['text'][:80]}")
                else:
                    print(f"  [{i}] TABLE: {item['rows']}x{item['cols']} - Headers: {item['header_hint'][:3]}")

            if len(result['structure']) > 30:
                print(f"  ... and {len(result['structure']) - 30} more elements")

        except Exception as e:
            print(f"Error processing {doc_path}: {e}")
            import traceback
            traceback.print_exc()


if __name__ == "__main__":
    main()
