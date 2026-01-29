"""
Explore one of the other documents to understand structure differences.
"""

from docx import Document
from pathlib import Path

def explore_table(table, table_index: int):
    """Extract detailed information from a table."""
    print(f"\n{'='*60}")
    print(f"TABLE {table_index}")
    print('='*60)

    rows = len(table.rows)
    cols = len(table.columns) if table.rows else 0
    print(f"Dimensions: {rows} rows x {cols} cols")

    # Extract all rows (limit to first 15 for display)
    display_rows = min(rows, 15)
    for i, row in enumerate(table.rows[:display_rows]):
        cells = [cell.text.strip().replace('\n', ' ')[:50] for cell in row.cells]
        print(f"  Row {i}: {cells}")

    if rows > display_rows:
        print(f"  ... {rows - display_rows} more rows")


def explore_document(doc_path: str):
    """Explore document structure."""
    doc = Document(doc_path)

    print(f"\nAnalyzing: {doc_path}")
    print(f"Total Paragraphs: {len(doc.paragraphs)}")
    print(f"Total Tables: {len(doc.tables)}")

    # Print all headings
    print("\n" + "="*60)
    print("DOCUMENT HEADINGS")
    print("="*60)
    for para in doc.paragraphs:
        if para.style and 'Heading' in para.style.name:
            print(f"  [{para.style.name}] {para.text[:80]}")

    # Explore each table
    for i, table in enumerate(doc.tables):
        explore_table(table, i)


def main():
    # Explore one of the other documents
    doc_path = Path("documents/Outlet_KYC _UB_3334.docx")
    explore_document(doc_path)

    print("\n\n" + "="*80)
    print("CHANGE BENEFICIARY DOCUMENT")
    print("="*80)
    doc_path2 = Path("documents/Change Beneficiary - UB 3526.docx")
    explore_document(doc_path2)


if __name__ == "__main__":
    main()
