"""
Deep exploration of the Vendor Creation template document.
"""

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from pathlib import Path
import json

def explore_table(table, table_index: int):
    """Extract detailed information from a table."""
    print(f"\n{'='*60}")
    print(f"TABLE {table_index}")
    print('='*60)

    rows = len(table.rows)
    cols = len(table.columns) if table.rows else 0
    print(f"Dimensions: {rows} rows x {cols} cols")

    # Extract all rows (limit to first 20 for display)
    display_rows = min(rows, 25)
    for i, row in enumerate(table.rows[:display_rows]):
        cells = [cell.text.strip().replace('\n', ' ')[:60] for cell in row.cells]
        print(f"  Row {i}: {cells}")

    if rows > display_rows:
        print(f"  ... {rows - display_rows} more rows")


def explore_full_document(doc_path: str):
    """Full exploration of document structure."""
    doc = Document(doc_path)

    print(f"\nAnalyzing: {doc_path}")
    print(f"Total Paragraphs: {len(doc.paragraphs)}")
    print(f"Total Tables: {len(doc.tables)}")

    # Print all headings
    print("\n" + "="*60)
    print("DOCUMENT HEADINGS")
    print("="*60)
    for para in doc.paragraphs:
        if para.style and 'Heading' in para.style.name:
            print(f"  [{para.style.name}] {para.text[:100]}")

    # Explore each table in detail
    for i, table in enumerate(doc.tables):
        explore_table(table, i)


def main():
    # Focus on the template document
    doc_path = Path("documents/Vendor Creation Sample BUD(1).docx")
    explore_full_document(doc_path)


if __name__ == "__main__":
    main()
