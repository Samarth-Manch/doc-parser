"""
section_parser.py
Generic BUD document parser for the pre-validator.

Parses configured sections from a .docx file into a single structured
representation (BUDDocument) that ALL validators consume.  No validator
should ever access the raw python-docx Document directly.

To add a new section:
  1. Add an entry to SECTION_HINTS.
  2. If the section contains field tables, add the key to FIELD_SECTIONS.
"""

import os
import re
import zipfile
from collections import defaultdict
from dataclasses import dataclass, field
from typing import Any

from docx import Document
from lxml import etree


# ── Configuration ────────────────────────────────────────────────────────────
# To parse a new section, add an entry here.
# hint[0] = descriptive heading text (lowercase), hint[1] = section number.

SECTION_HINTS: dict[str, list[str]] = {
    "4.4":   ["4.4 field-level information", "4.4"],
    "4.5.1": ["4.5.1", "4.5.1"],
    "4.5.2": ["4.5.2", "4.5.2"],
    "4.6":   ["4.6 reference tables", "4.6"],
    "4.7":   ["4.7 record list view", "4.7"],
}

# Sections that contain field tables (name, type columns at minimum).
FIELD_SECTIONS: set[str] = {"4.4", "4.5.1", "4.5.2"}

# Column name candidates (consolidated from all validators).
NAME_CANDIDATES = ("field name", "filed name", "name")
TYPE_CANDIDATES = ("field type", "fieldtype", "type", "data type")
MANDATORY_CANDIDATES = ("mandatory", "required", "is mandatory")
LOGIC_CANDIDATES = (
    "logic", "logic and rules", "logic and rules other than common logic",
    "rules", "rule", "validation", "description",
)

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_NEXT_SECTION_PAT = re.compile(r"^\d+\.\d+")


# ── Data structures ─────────────────────────────────────────────────────────

@dataclass
class FieldRow:
    """A single field extracted from a field table."""
    name: str               # original case, stripped
    name_lower: str         # lowercased for matching
    field_type: str         # original case, stripped
    field_type_upper: str   # uppercased for matching
    is_mandatory: bool
    logic: str              # full logic text
    rules: str              # separate "rules" column if present
    panel: str              # panel name (original case), "" if before first PANEL
    row_index: int          # 1-based row in the table (header = row 1)
    table_index: int        # 0-based index of table within the section


@dataclass
class TableData:
    """Raw table extracted from a section."""
    headers: list[str]           # lowercased
    headers_original: list[str]  # original case
    rows: list[list[str]]        # data rows (original case, stripped)
    is_field_table: bool         # has at least name + type columns
    column_map: dict[str, int]   # matched canonical name -> column index


@dataclass
class SectionData:
    """All parsed data for a single section."""
    key: str
    heading_found: bool
    heading_text: str = ""
    heading_para_idx: int | None = None
    is_not_applicable: bool = False
    raw_tables: list[TableData] = field(default_factory=list)
    field_rows: list[FieldRow] = field(default_factory=list)
    fields_by_name: dict[str, FieldRow] = field(default_factory=dict)
    panels: dict[str, list[FieldRow]] = field(default_factory=dict)
    # Paragraphs between heading and next heading (plain text, for section-text checks)
    section_text: str = ""

    @property
    def fields_dict(self) -> dict[str, dict]:
        """Fields as {name_lower: {"type": ..., "logic": ..., "is_mandatory": ...}}.

        Excludes PANEL rows.  First occurrence wins for duplicate names.
        """
        result: dict[str, dict] = {}
        for f in self.field_rows:
            if f.field_type_upper == "PANEL":
                continue
            if f.name_lower not in result:
                result[f.name_lower] = {
                    "type": f.field_type,
                    "logic": f.logic,
                    "is_mandatory": f.is_mandatory,
                }
        return result

    @property
    def has_fields(self) -> bool:
        """True if this section has any non-PANEL field rows."""
        return any(f.field_type_upper != "PANEL" for f in self.field_rows)


@dataclass
class BUDDocument:
    """Complete parsed BUD document — the single source of truth for all validators."""
    sections: dict[str, SectionData]
    docx_path: str
    ole_label_map: dict[str, str] = field(default_factory=dict)
    section_46_table_names: set[str] = field(default_factory=set)
    section_47_fields: list[str] = field(default_factory=list)


# ── Column helpers ───────────────────────────────────────────────────────────

def _get_col_index(headers: list[str], *candidates: str) -> int | None:
    """Return the index of the first matching column name, or None."""
    for name in candidates:
        if name in headers:
            return headers.index(name)
    return None


def _is_mandatory_value(text: str) -> bool:
    return text.strip().lower() in ("yes", "true", "y", "1", "mandatory")


# ── Heading finders ──────────────────────────────────────────────────────────

def _heading_level(style_name: str) -> int:
    m = re.search(r"(\d+)", style_name)
    return int(m.group(1)) if m else 99


def _find_section_heading(doc: Document, section_key: str) -> tuple[int | None, str]:
    """Find paragraph index and text for a section heading.

    Strategy (most specific to most lenient):
      1. Heading-styled paragraph matching descriptive hint
      2. Heading-styled paragraph matching section number
      3. Any paragraph containing descriptive hint text
      4. Any paragraph starting with section number
    """
    hints = SECTION_HINTS[section_key]
    desc_hint = hints[0].lower()
    num_hint = hints[1] if len(hints) > 1 else section_key

    # Pass 1: Heading-styled, descriptive
    for i, para in enumerate(doc.paragraphs):
        if para.style.name.startswith("Heading") and desc_hint in para.text.strip().lower():
            return i, para.text.strip()

    # Pass 2: Heading-styled, section number
    for i, para in enumerate(doc.paragraphs):
        if para.style.name.startswith("Heading") and num_hint in para.text:
            return i, para.text.strip()

    # Pass 3: Any paragraph containing descriptive text
    for i, para in enumerate(doc.paragraphs):
        if desc_hint in para.text.strip().lower():
            return i, para.text.strip()

    # Pass 4: Any paragraph starting with section number
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith(section_key):
            return i, para.text.strip()

    return None, ""


# ── Body position helpers ────────────────────────────────────────────────────

def _body_index(doc: Document, element) -> int:
    for i, elem in enumerate(doc.element.body):
        if elem is element:
            return i
    return -1


def _section_end_pos(doc: Document, heading_para_idx: int) -> int:
    """Body-element index where this section ends.

    Stops at next heading of same/higher level OR a numbered-section heading.
    """
    level = _heading_level(doc.paragraphs[heading_para_idx].style.name)
    body = list(doc.element.body)

    for para in doc.paragraphs[heading_para_idx + 1:]:
        if not para.style.name.startswith("Heading"):
            continue
        lvl = _heading_level(para.style.name)
        text = para.text.strip()
        if lvl <= level or _NEXT_SECTION_PAT.match(text):
            return body.index(para._element)

    return len(body)


def _tables_between(doc: Document, start_pos: int, end_pos: int) -> list:
    """Return python-docx table objects between two body positions."""
    result = []
    for table in doc.tables:
        pos = _body_index(doc, table._tbl)
        if start_pos < pos < end_pos:
            result.append(table)
    return result


# ── Table extraction ─────────────────────────────────────────────────────────

def _extract_table_data(table) -> TableData:
    """Convert a python-docx table into a TableData object."""
    if not table.rows:
        return TableData([], [], [], False, {})

    headers_orig = [c.text.strip() for c in table.rows[0].cells]
    headers_lower = [h.lower() for h in headers_orig]

    rows = []
    for row in table.rows[1:]:
        rows.append([c.text.strip() for c in row.cells])

    # Check if this is a field table (has name + type columns)
    name_idx = _get_col_index(headers_lower, *NAME_CANDIDATES)
    type_idx = _get_col_index(headers_lower, *TYPE_CANDIDATES)
    is_field_table = (
        name_idx is not None
        and type_idx is not None
        and len(table.rows) >= 2
    )

    # Build column map
    col_map: dict[str, int] = {}
    if name_idx is not None:
        col_map["name"] = name_idx
    if type_idx is not None:
        col_map["type"] = type_idx
    mand_idx = _get_col_index(headers_lower, *MANDATORY_CANDIDATES)
    if mand_idx is not None:
        col_map["mandatory"] = mand_idx
    logic_idx = _get_col_index(headers_lower, *LOGIC_CANDIDATES)
    if logic_idx is not None:
        col_map["logic"] = logic_idx
    # Separate "rules" column (if distinct from matched logic column)
    rules_idx = _get_col_index(headers_lower, "rules", "rule")
    if rules_idx is not None and rules_idx != logic_idx:
        col_map["rules"] = rules_idx

    return TableData(
        headers=headers_lower,
        headers_original=headers_orig,
        rows=rows,
        is_field_table=is_field_table,
        column_map=col_map,
    )


# ── Field extraction (for field sections) ────────────────────────────────────

def _extract_field_rows(
    tables: list[TableData],
) -> tuple[list[FieldRow], dict[str, FieldRow], dict[str, list[FieldRow]]]:
    """Extract FieldRow objects from field tables, tracking panels.

    Returns (field_rows, fields_by_name, panels).
    - field_rows: ALL rows in order, including PANEL rows and duplicates
    - fields_by_name: non-PANEL fields, first occurrence, lowercase key
    - panels: non-PANEL fields grouped by panel name
    """
    field_rows: list[FieldRow] = []
    fields_by_name: dict[str, FieldRow] = {}
    panels: dict[str, list[FieldRow]] = defaultdict(list)
    current_panel = ""

    for t_idx, table in enumerate(tables):
        if not table.is_field_table:
            continue

        name_idx = table.column_map.get("name")
        type_idx = table.column_map.get("type")
        mand_idx = table.column_map.get("mandatory")
        logic_idx = table.column_map.get("logic")
        rules_idx = table.column_map.get("rules")

        if name_idx is None or type_idx is None:
            continue

        for r_idx, row in enumerate(table.rows):
            fname = row[name_idx] if name_idx < len(row) else ""
            ftype = row[type_idx] if type_idx < len(row) else ""

            if not fname.strip():
                continue

            is_panel = ftype.strip().upper() == "PANEL"
            if is_panel:
                current_panel = fname.strip()

            is_mand = False
            if mand_idx is not None and mand_idx < len(row):
                is_mand = _is_mandatory_value(row[mand_idx])

            logic = ""
            if logic_idx is not None and logic_idx < len(row):
                logic = row[logic_idx]

            rules_text = ""
            if rules_idx is not None and rules_idx < len(row):
                rules_text = row[rules_idx]

            fr = FieldRow(
                name=fname.strip(),
                name_lower=fname.strip().lower(),
                field_type=ftype.strip(),
                field_type_upper=ftype.strip().upper(),
                is_mandatory=is_mand,
                logic=logic,
                rules=rules_text,
                panel=current_panel if not is_panel else fname.strip(),
                row_index=r_idx + 2,  # 1-based; row 1 = header
                table_index=t_idx,
            )
            field_rows.append(fr)

            if not is_panel:
                if fr.name_lower not in fields_by_name:
                    fields_by_name[fr.name_lower] = fr
                panels[current_panel].append(fr)

    return field_rows, fields_by_name, dict(panels)


# ── Section 4.6 — reference table name extraction ───────────────────────────

_NON_TABLE_KEYWORDS = frozenset({
    "edv", "staging", "data", "table", "of", "the", "reference", "refer",
})

# Trailing words in 4.6 entries that are often omitted in logic references
_TRAILING_SUFFIXES = re.compile(
    r"[_\s]+(details|data|info|list|master)$", re.IGNORECASE,
)


def normalise_table_name(name: str) -> str:
    """Lower-case, strip leading 'table' prefix, unify separators."""
    n = name.strip().lower()
    n = re.sub(r"^table\s+", "", n)
    n = re.sub(r"[-_\s]+", "_", n)
    return n


def normalise_table_name_variants(name: str) -> list[str]:
    """Return the normalised name plus a variant with common trailing suffixes stripped."""
    base = normalise_table_name(name)
    variants = [base]
    stripped = _TRAILING_SUFFIXES.sub("", base)
    if stripped and stripped != base:
        variants.append(stripped)
    return variants


def _extract_table_identifiers(text: str) -> list[str]:
    """Parse a single text entry from section 4.6 and return table identifiers."""
    ids: list[str] = []
    t = text.strip()
    if not t:
        return ids

    # Numbered table  (e.g. "Table 1.9 - PAYMENT_MODE")
    num_match = re.search(
        r"(?:(?:refer(?:ence)?\s+)?table\s+)?(\d+\.\d+)", t, re.IGNORECASE,
    )
    if num_match:
        ids.append(num_match.group(1))

    # UPPER_CASE identifiers with underscores
    for name in re.findall(r"\b[A-Za-z_][A-Za-z0-9]*(?:_[A-Za-z0-9_]+)+\b", t):
        if name.lower() not in _NON_TABLE_KEYWORDS:
            ids.append(name)

    # Parenthesized names
    for name in re.findall(r"\(([A-Za-z_][A-Za-z0-9_]*)\)", t):
        if name.lower() not in _NON_TABLE_KEYWORDS:
            ids.append(name)

    # Single ALL-CAPS identifier
    if not ids and re.match(r"^[A-Z][A-Z0-9]+$", t) and t.lower() not in _NON_TABLE_KEYWORDS:
        ids.append(t)

    # Short plain text entries (≤6 words, no sentence punctuation)
    if not ids and not re.search(r"[.;:!?]", t) and len(t.split()) <= 6:
        cleaned = t.strip()
        if cleaned and cleaned.lower() not in _NON_TABLE_KEYWORDS:
            ids.append(cleaned)

    return ids


def _extract_46_table_names(
    doc: Document,
    section: SectionData,
    docx_path: str,
    ole_map: dict[str, str],
) -> set[str]:
    """Extract normalised reference-table identifiers from section 4.6."""
    names: set[str] = set()

    if section.heading_para_idx is None:
        return names

    body = list(doc.element.body)
    start = _body_index(doc, doc.paragraphs[section.heading_para_idx]._element)
    end = _section_end_pos(doc, section.heading_para_idx)

    for idx, elem in enumerate(body):
        if idx <= start or idx >= end:
            continue

        tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag

        if tag == "p":
            # Visible text
            visible_parts = [
                t_elem.text
                for t_elem in elem.iter(f"{{{_W_NS}}}t")
                if t_elem.text
            ]
            visible = "".join(visible_parts).strip()
            if visible:
                for ident in _extract_table_identifiers(visible):
                    names.update(normalise_table_name_variants(ident))

            # OLE LINK field codes (embedded Excel file paths)
            xml_str = etree.tostring(elem, encoding="unicode")
            for m in re.finditer(r'"([^"]+\.xlsx?)"', xml_str, re.IGNORECASE):
                fname = os.path.basename(m.group(1))
                name_no_ext = os.path.splitext(fname)[0]
                names.update(normalise_table_name_variants(name_no_ext))

        elif tag == "tbl":
            for t_elem in elem.iter(f"{{{_W_NS}}}t"):
                if t_elem.text:
                    for ident in _extract_table_identifiers(t_elem.text.strip()):
                        names.update(normalise_table_name_variants(ident))

    # OLE label map entries
    for _filename, label in ole_map.items():
        for ident in _extract_table_identifiers(label):
            names.update(normalise_table_name_variants(ident))

    return names


# ── Section 4.7 — record list view field extraction ─────────────────────────

def _extract_47_fields(doc: Document, section: SectionData) -> list[str]:
    """Extract bullet-point field names from section 4.7."""
    if section.heading_para_idx is None:
        return []

    heading_elem = doc.paragraphs[section.heading_para_idx]._element
    start_pos = _body_index(doc, heading_elem)
    end_pos = _section_end_pos(doc, section.heading_para_idx)

    fields: list[str] = []
    desc_keywords = ["list ", "below", "columns", "displayed", "dashboards", "following"]

    for para in doc.paragraphs[section.heading_para_idx + 1:]:
        pos = _body_index(doc, para._element)
        if pos >= end_pos:
            break

        text = para.text.strip()
        if not text:
            continue

        if para.style.name.startswith("Heading"):
            break

        # Skip descriptive paragraphs
        lower = text.lower()
        if "\n" in text or len(text) > 120 or any(kw in lower for kw in desc_keywords):
            continue

        # Strip bullet prefixes
        for prefix in ("\u2022", "-", "\u2013", "\u2014", "*"):
            if text.startswith(prefix):
                text = text[len(prefix):].strip()
                break

        text = text.rstrip(".;,")
        if text:
            fields.append(text)

    return fields


# ── OLE label map ────────────────────────────────────────────────────────────

def _build_ole_label_map(docx_path: str) -> dict[str, str]:
    """Map embedded OLE filenames to their preceding paragraph text."""
    label_map: dict[str, str] = {}
    try:
        with zipfile.ZipFile(docx_path, "r") as zf:
            # Step 1: rId -> filename from relationships
            rid_to_file: dict[str, str] = {}
            try:
                rels_xml = zf.read("word/_rels/document.xml.rels")
                rels_tree = etree.fromstring(rels_xml)
                for rel in rels_tree:
                    target = rel.get("Target", "")
                    if "embeddings/" in target.lower():
                        rid_to_file[rel.get("Id")] = target.split("/")[-1]
            except (KeyError, etree.XMLSyntaxError):
                return label_map

            if not rid_to_file:
                return label_map

            # Step 2: walk body paragraphs, track last text, detect OLE refs
            doc_xml = zf.read("word/document.xml")
            tree = etree.fromstring(doc_xml)

            w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            body = tree.find(f"{{{w_ns}}}body")
            if body is None:
                return label_map

            last_text = ""
            for elem in body:
                if etree.QName(elem).localname != "p":
                    continue
                text = "".join(elem.itertext()).strip()
                if text:
                    last_text = text
                elem_str = etree.tostring(elem, encoding="unicode")
                for rid, filename in rid_to_file.items():
                    if f'r:id="{rid}"' in elem_str and last_text:
                        if filename not in label_map:
                            label_map[filename] = last_text
    except Exception:
        pass
    return label_map


# ── Section text extraction ──────────────────────────────────────────────────

def _extract_section_text(doc: Document, heading_para_idx: int, end_pos: int) -> str:
    """Collect plain text from paragraphs between heading and section end."""
    parts: list[str] = []
    for para in doc.paragraphs[heading_para_idx + 1:]:
        if _body_index(doc, para._element) >= end_pos:
            break
        parts.append(para.text.strip())
    return " ".join(parts)


# ── Main entry point ─────────────────────────────────────────────────────────

def parse_bud(docx_path: str) -> BUDDocument:
    """Parse a BUD .docx file into a structured BUDDocument.

    This is the single entry point.  After calling this, no validator needs
    to access the raw python-docx Document.
    """
    doc = Document(docx_path)
    sections: dict[str, SectionData] = {}

    for section_key in SECTION_HINTS:
        para_idx, heading_text = _find_section_heading(doc, section_key)

        sd = SectionData(
            key=section_key,
            heading_found=para_idx is not None,
            heading_text=heading_text,
            heading_para_idx=para_idx,
        )

        if para_idx is not None:
            start_pos = _body_index(doc, doc.paragraphs[para_idx]._element)
            end_pos = _section_end_pos(doc, para_idx)

            # Extract raw tables
            raw_docx_tables = _tables_between(doc, start_pos, end_pos)
            sd.raw_tables = [_extract_table_data(t) for t in raw_docx_tables]

            # Section text (for "Not Applicable" detection and other checks)
            sd.section_text = _extract_section_text(doc, para_idx, end_pos)

            # "Not Applicable" detection
            if not raw_docx_tables:
                sd.is_not_applicable = "not applicable" in sd.section_text.lower()

            # Extract structured fields for field sections
            if section_key in FIELD_SECTIONS:
                sd.field_rows, sd.fields_by_name, sd.panels = _extract_field_rows(
                    sd.raw_tables,
                )

        sections[section_key] = sd

    # OLE label map
    ole_map = _build_ole_label_map(docx_path)

    # Section 4.6 reference table names
    section_46_names: set[str] = set()
    s46 = sections.get("4.6")
    if s46 and s46.heading_found:
        section_46_names = _extract_46_table_names(doc, s46, docx_path, ole_map)

    # Section 4.7 bullet fields
    section_47_fields: list[str] = []
    s47 = sections.get("4.7")
    if s47 and s47.heading_found:
        section_47_fields = _extract_47_fields(doc, s47)

    return BUDDocument(
        sections=sections,
        docx_path=docx_path,
        ole_label_map=ole_map,
        section_46_table_names=section_46_names,
        section_47_fields=section_47_fields,
    )
