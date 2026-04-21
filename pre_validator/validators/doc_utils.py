"""
doc_utils.py
Shared document-traversal helpers used by multiple validators.

All functions operate on a python-docx Document object.
"""

import re
import zipfile
from collections import defaultdict
from typing import Iterator

from lxml import etree
from docx import Document


# Section numbers mapped to heading hints for lookup
SECTION_HINTS = {
    "4.4":   ["4.4 field-level information", "4.4"],
    "4.5.1": ["4.5.1", "4.5.1"],
    "4.5.2": ["4.5.2", "4.5.2"],
    "4.7":   ["4.7 record list view", "4.7"],
}

REQUIRED_SECTIONS = [
    ("4.4",   "4.4 Field-Level Information"),
    ("4.5.1", "4.5.1 Initiator Behaviour"),
    ("4.5.2", "4.5.2 Vendor Behaviour"),
]

FIELD_TABLE_HEADERS = {"field name", "field type", "mandatory", "logic"}


# ── Element position helpers ────────────────────────────────────────────────

def body_index(doc: Document, element) -> int:
    """Return the index of *element* inside doc.element.body, or -1."""
    for i, elem in enumerate(doc.element.body):
        if elem is element:
            return i
    return -1


def paragraph_position(doc: Document, para_idx: int) -> int:
    """Return the body-element index for a given paragraph index."""
    return body_index(doc, doc.paragraphs[para_idx]._element)


# ── Section heading finders ─────────────────────────────────────────────────

def find_section_heading_index(doc: Document, heading_text: str) -> int | None:
    """Return the paragraph index whose text contains *heading_text*."""
    for i, para in enumerate(doc.paragraphs):
        if heading_text in para.text:
            return i
    return None


def find_heading_by_number(doc: Document, section_num: str) -> int | None:
    """Fallback: first paragraph whose text starts with *section_num*."""
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith(section_num):
            return i
    return None


def find_section_heading(doc: Document, section_key: str) -> int | None:
    """
    Find a section heading by its key (e.g. '4.4', '4.5.1').

    Tries a descriptive match first, then falls back to section number.
    Only considers paragraphs styled as Heading*.
    """
    hints = SECTION_HINTS[section_key]
    # Descriptive match
    for i, para in enumerate(doc.paragraphs):
        if not para.style.name.startswith("Heading"):
            continue
        if hints[0].lower() in para.text.strip().lower():
            return i
    # Fallback: section number in heading
    for i, para in enumerate(doc.paragraphs):
        if not para.style.name.startswith("Heading"):
            continue
        if hints[1] in para.text:
            return i
    return None


# ── Section boundary helpers ────────────────────────────────────────────────

def next_heading_position(doc: Document, heading_para_idx: int) -> int:
    """
    Return body-element index of the next Heading after *heading_para_idx*,
    or the total body length if there is none.
    """
    body = list(doc.element.body)
    for para in doc.paragraphs[heading_para_idx + 1:]:
        if para.style.name.startswith("Heading"):
            return body.index(para._element)
    return len(body)


def tables_between(doc: Document, start_pos: int, end_pos: int) -> list:
    """Return all tables whose body position falls strictly between start and end."""
    result = []
    for table in doc.tables:
        pos = body_index(doc, table._tbl)
        if start_pos < pos < end_pos:
            result.append(table)
    return result


def tables_under_section(doc: Document, heading_para_idx: int) -> list:
    """Return all tables between the heading and the next heading."""
    start = body_index(doc, doc.paragraphs[heading_para_idx]._element)
    end = next_heading_position(doc, heading_para_idx)
    return tables_between(doc, start, end)


def tables_in_section(doc: Document, heading_text: str) -> list:
    """Return all tables between the named heading and the next heading."""
    heading_idx = find_section_heading_index(doc, heading_text)
    if heading_idx is None:
        return []
    start = paragraph_position(doc, heading_idx)
    end = next_heading_position(doc, heading_idx)
    return tables_between(doc, start, end)


# ── Table type detection ────────────────────────────────────────────────────

def is_field_table(table) -> bool:
    """Return True if the table's header row matches field-table column names."""
    if not table.rows:
        return False
    header_cells = {cell.text.strip().lower() for cell in table.rows[0].cells}
    return bool(FIELD_TABLE_HEADERS & header_cells)


def is_field_table_strict(table) -> bool:
    """Return True if the table has both a name column and a type column."""
    if len(table.rows) < 2:
        return False
    headers = [c.text.strip().lower() for c in table.rows[0].cells]
    has_name = "field name" in headers or "filed name" in headers
    has_type = "field type" in headers or "type" in headers
    return has_name and has_type


def get_column_index(headers: list[str], *candidates: str) -> int | None:
    """Return the index of the first matching column name, or None."""
    for col_name in candidates:
        if col_name in headers:
            return headers.index(col_name)
    return None


# ── Shared field helpers ───────────────────────────────────────────────────

def group_fields_by_panel(fields: list) -> dict[str, list]:
    """Group FieldDefinition objects by their panel (section attribute).
    Skips PANEL-type rows themselves."""
    panels: dict[str, list] = defaultdict(list)
    for f in fields:
        panel_name = (f.section or "").strip() or "(no panel)"
        if hasattr(f, "field_type_raw") and f.field_type_raw and f.field_type_raw.strip().upper() == "PANEL":
            continue
        panels[panel_name].append(f)
    return panels


def iter_field_tables(doc: Document) -> Iterator[tuple[str, list]]:
    """Yield (section_key, field_tables) for each section that has field tables."""
    for section_key in SECTION_HINTS:
        heading_idx = find_section_heading(doc, section_key)
        if heading_idx is None:
            continue
        tables = tables_under_section(doc, heading_idx)
        field_tbls = [t for t in tables if is_field_table_strict(t)]
        if field_tbls:
            yield section_key, field_tbls


# ── OLE embedded Excel label mapping ──────────────────────────────────────

def build_ole_label_map(docx_path: str) -> dict[str, str]:
    """
    Map embedded OLE filenames to their preceding bullet/paragraph text.

    LibreOffice-embedded Excel files get generic names (oleObject1.xlsx, etc.)
    unlike MS Word which preserves meaningful names. However the BUD document
    has bulleted text above each embedded object with the proper table name
    (e.g. "Table 1.1", "Vendor Master").

    Walks the DOCX XML body in document order: for each paragraph that
    references an OLE embedding (via r:id), records the last non-empty
    paragraph text seen before it as the label.

    Returns:
        dict mapping embedded filename (e.g. "oleObject1.xlsx") to
        the preceding paragraph text (e.g. "Table 1.1").
        Empty dict if no mappings could be built.
    """
    label_map: dict[str, str] = {}

    try:
        with zipfile.ZipFile(docx_path, "r") as zf:
            # Step 1: rId -> filename from relationships
            rid_to_file: dict[str, str] = {}
            try:
                rels_xml = zf.read("word/_rels/document.xml.rels")
                rels_tree = etree.fromstring(rels_xml)
                for rel in rels_tree:
                    target = rel.get("Target", "")
                    if "embeddings/" in target.lower():
                        rid_to_file[rel.get("Id")] = target.split("/")[-1]
            except (KeyError, etree.XMLSyntaxError):
                return label_map

            if not rid_to_file:
                return label_map

            # Step 2: walk body paragraphs, track last text, detect OLE refs
            doc_xml = zf.read("word/document.xml")
            tree = etree.fromstring(doc_xml)

            w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            body = tree.find(f"{{{w_ns}}}body")
            if body is None:
                return label_map

            last_text = ""
            for elem in body:
                if etree.QName(elem).localname != "p":
                    continue

                text = "".join(elem.itertext()).strip()
                if text:
                    last_text = text

                # Check for rId references pointing to embedded files
                elem_str = etree.tostring(elem, encoding="unicode")
                for rid, filename in rid_to_file.items():
                    if f'r:id="{rid}"' in elem_str and last_text:
                        if filename not in label_map:
                            label_map[filename] = last_text

    except (zipfile.BadZipFile, FileNotFoundError, Exception):
        pass

    return label_map
