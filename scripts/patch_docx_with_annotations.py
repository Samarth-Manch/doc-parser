#!/usr/bin/env python3
"""Patch a BUD .docx by inserting cross-panel annotations produced by
annotate_cross_panel_refs.py.

Annotations are always INSERT-only additions of the form ` (from X panel)`.
The annotator operates on whitespace-normalized text (paragraphs joined with
space by the parser) while the raw docx has its own paragraph/run structure.

To bridge that, for each changed field we:
  1. Diff old vs new to extract a list of (preceding_anchor, insert_text) pairs
     where anchor = the text from old up to the insertion point.
  2. For every table cell, check if its whitespace-normalized text equals the
     normalized `old`. If so, for each insertion, locate the anchor within the
     cell using whitespace-insensitive search, then walk the paragraph runs to
     splice the insertion at the right run-offset.
"""
import argparse
import difflib
import json
import re
import shutil
from pathlib import Path

from docx import Document


WS_RE = re.compile(r"\s+")


def norm(s: str) -> str:
    return WS_RE.sub(" ", s).strip()


def compute_insertions(old: str, new: str):
    """Return list of dicts: {anchor_in_old, insert_text, old_pos}.

    anchor_in_old = slice of old[0:old_pos] (everything before the insertion).
    """
    sm = difflib.SequenceMatcher(a=old, b=new, autojunk=False)
    insertions = []
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "insert" and j1 < j2:
            insertions.append({
                "anchor_in_old": old[:i1],
                "insert_text": new[j1:j2],
                "old_pos": i1,
                "pure_insert": True,
            })
        elif tag in ("replace", "delete") and j1 < j2:
            insertions.append({
                "anchor_in_old": old[:i1],
                "insert_text": new[j1:j2],
                "old_pos": i1,
                "pure_insert": False,
            })
    return insertions


def build_ws_index(raw: str):
    """Return (normalized_text, mapping) where mapping[k] = raw_index of the
    k-th char in normalized_text. Also returns len_after for end-mapping.
    """
    out_chars = []
    mapping = []
    last_was_space = True  # so leading whitespace is trimmed
    i = 0
    for ch in raw:
        if ch.isspace() or ch == "\n" or ch == "\t":
            if not last_was_space:
                out_chars.append(" ")
                mapping.append(i)
                last_was_space = True
        else:
            out_chars.append(ch)
            mapping.append(i)
            last_was_space = False
        i += 1
    # strip trailing space in normalized
    while out_chars and out_chars[-1] == " ":
        out_chars.pop()
        mapping.pop()
    # for end-mapping past the last char, we also need to know raw length
    return "".join(out_chars), mapping, len(raw)


def ws_find(normalized: str, needle_normalized: str, start: int = 0) -> int:
    """Locate needle within normalized text. Both sides already whitespace-normalized."""
    return normalized.find(needle_normalized, start)


def cell_concat_paragraphs(cell):
    """Return list of (paragraph_index, paragraph, start_offset_in_raw_cell_text).

    cell text is computed as '\\n'.join(p.text) exactly (matches cell.text logic).
    """
    positions = []
    running = 0
    for i, para in enumerate(cell.paragraphs):
        positions.append((i, para, running))
        running += len(para.text) + 1  # +1 for the inter-paragraph \n
    # strip trailing \n for the last paragraph
    total_len = max(0, running - 1)
    return positions, total_len


def raw_cell_text(cell) -> str:
    return "\n".join(p.text for p in cell.paragraphs)


def insert_text_into_paragraph(paragraph, char_offset: int, text: str) -> bool:
    """Insert text at char_offset within the paragraph by splicing a run."""
    runs = paragraph.runs
    pos = 0
    for run in runs:
        run_text = run.text or ""
        run_len = len(run_text)
        if pos <= char_offset <= pos + run_len:
            local = char_offset - pos
            run.text = run_text[:local] + text + run_text[local:]
            return True
        pos += run_len
    if runs:
        # past last run: append to last
        runs[-1].text = (runs[-1].text or "") + text
        return True
    if not runs and char_offset == 0:
        paragraph.add_run(text)
        return True
    return False


def raw_offset_to_paragraph(positions, raw_offset: int):
    """Given [(i, para, start_offset), ...] return (paragraph, local_offset).

    raw_offset refers to a position within '\\n'.join(p.text for p in cell.paragraphs).
    When raw_offset falls on an inter-paragraph '\\n', we bind it to the end of
    the earlier paragraph.
    """
    for idx, (i, para, start) in enumerate(positions):
        para_len = len(para.text)
        end = start + para_len  # exclusive of newline
        if start <= raw_offset <= end:
            return para, raw_offset - start
        # if raw_offset == end + 1 it's past the newline, i.e. start of next
    # past end: stick to last paragraph
    if positions:
        i, para, start = positions[-1]
        return para, len(para.text)
    return None, None


def find_cell_with_logic(doc, old_logic_norm: str):
    """Return (cell, raw_cell_text, normalized_text, mapping, raw_len) or None."""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                raw = raw_cell_text(cell)
                normed, mapping, raw_len = build_ws_index(raw)
                if old_logic_norm and old_logic_norm in normed:
                    return cell, raw, normed, mapping, raw_len
    return None


def apply_insertions(doc, old: str, new: str, insertions) -> dict:
    old_norm = norm(old)
    found = find_cell_with_logic(doc, old_norm)
    if not found:
        return {"ok": False, "reason": "cell not found (normalized match failed)"}
    cell, raw, normed, mapping, raw_len = found
    # offset of old within normed
    old_start_in_norm = normed.find(old_norm)
    positions, _ = cell_concat_paragraphs(cell)

    # Apply from last to first so earlier insertions don't shift later raw offsets.
    ordered = sorted(insertions, key=lambda x: x["old_pos"], reverse=True)
    results = []
    for ins in ordered:
        anchor_norm = norm(ins["anchor_in_old"])
        # anchor may be empty (insertion at start) — then insert at beginning of old's location
        if anchor_norm:
            anchor_pos_in_norm = normed.find(anchor_norm, old_start_in_norm)
            if anchor_pos_in_norm == -1:
                anchor_pos_in_norm = normed.find(anchor_norm)
            if anchor_pos_in_norm == -1:
                results.append({"ok": False, "reason": "anchor not found", "insert": ins["insert_text"]})
                continue
            # Place insertion just before the next normalized char after the anchor.
            # That skips over any raw whitespace between anchor's last char and the next word,
            # so an insert like " (from X panel) " slots in cleanly without double-spacing.
            end_pos = anchor_pos_in_norm + len(anchor_norm)
            if end_pos < len(mapping):
                raw_offset = mapping[end_pos]
            else:
                raw_offset = raw_len
        else:
            raw_offset = mapping[old_start_in_norm] if old_start_in_norm < len(mapping) else 0

        para, local = raw_offset_to_paragraph(positions, raw_offset)
        if para is None:
            results.append({"ok": False, "reason": "no paragraph", "insert": ins["insert_text"]})
            continue
        # Normalize whitespace around the insert to avoid doubled / missing spaces.
        core = ins["insert_text"].strip()
        left_ws = raw_offset > 0 and raw_offset <= len(raw) and raw[raw_offset - 1].isspace()
        right_ws = raw_offset < len(raw) and raw[raw_offset].isspace()
        spaced = ("" if left_ws else " ") + core + ("" if right_ws else " ")
        ok = insert_text_into_paragraph(para, local, spaced)
        results.append({"ok": ok, "insert": spaced, "raw_offset": raw_offset})

    return {"ok": all(r["ok"] for r in results), "results": results, "cell_preview": raw[:100]}


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--src", required=True)
    ap.add_argument("--dst", required=True)
    ap.add_argument("--annotations", required=True)
    ap.add_argument("--report", required=True)
    args = ap.parse_args()

    src = Path(args.src)
    dst = Path(args.dst)
    shutil.copyfile(src, dst)

    with open(args.annotations) as fh:
        annotations = json.load(fh)

    doc = Document(str(dst))
    report = {"applied": [], "failed": [], "skipped": []}

    changed = {fid: a for fid, a in annotations.items() if a.get("changed") and a.get("old") != a.get("new")}
    print(f"Annotations: {len(annotations)} total, {len(changed)} with changes")

    for fid, ann in changed.items():
        old = ann["old"]
        new = ann["new"]
        insertions = compute_insertions(old, new)
        if not insertions:
            report["skipped"].append({"field": fid, "reason": "no diff insertions"})
            continue
        outcome = apply_insertions(doc, old, new, insertions)
        if outcome["ok"]:
            report["applied"].append({"field": fid, "count": len(insertions)})
            print(f"[OK]   {fid}: {len(insertions)} insertion(s)")
        else:
            report["failed"].append({"field": fid, "reason": outcome.get("reason"), "results": outcome.get("results")})
            print(f"[FAIL] {fid}: {outcome.get('reason') or outcome.get('results')}")

    doc.save(str(dst))

    with open(args.report, "w") as fh:
        json.dump(report, fh, indent=2, ensure_ascii=False)

    print(f"\nSaved: {dst}")
    print(f"Applied: {len(report['applied'])}, Failed: {len(report['failed'])}, Skipped: {len(report['skipped'])}")


if __name__ == "__main__":
    main()
