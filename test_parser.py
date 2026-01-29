"""
Comprehensive test suite for document parser
Tests extraction accuracy by comparing text content with parsed JSON
"""

import unittest
import json
from pathlib import Path
from docx import Document
from doc_parser.parser import DocumentParser
from doc_parser.models import FieldType


class TextExtractor:
    """Extract all text content from document for comparison"""

    @staticmethod
    def extract_all_text(doc_path):
        """Extract all text from document including tables"""
        doc = Document(doc_path)

        text_content = {
            'paragraphs': [],
            'table_cells': [],
            'all_text': []
        }

        # Extract paragraph text
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                text_content['paragraphs'].append(text)
                text_content['all_text'].append(text)

        # Extract table text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        text_content['table_cells'].append(text)
                        text_content['all_text'].append(text)

        # Create searchable text
        text_content['full_text'] = ' '.join(text_content['all_text']).lower()

        return text_content

    @staticmethod
    def extract_field_names_from_tables(doc_path):
        """Extract all field names from tables"""
        doc = Document(doc_path)
        field_names = []

        for table in doc.tables:
            if not table.rows:
                continue

            # Check if this is a field definition table
            headers = [cell.text.strip().lower() for cell in table.rows[0].cells]

            if any('field name' in h or 'filed name' in h for h in headers):
                # This is a field table
                name_col = None
                for i, h in enumerate(headers):
                    if 'field name' in h or 'filed name' in h:
                        name_col = i
                        break

                if name_col is not None:
                    for row in table.rows[1:]:  # Skip header
                        if name_col < len(row.cells):
                            field_name = row.cells[name_col].text.strip()
                            if field_name:
                                field_names.append(field_name)

        return field_names


class TestDocumentParser(unittest.TestCase):
    """Test suite for document parser"""

    @classmethod
    def setUpClass(cls):
        """Setup test environment"""
        cls.parser = DocumentParser()
        cls.test_docs = list(Path("documents").glob("*.docx"))

        if not cls.test_docs:
            raise unittest.SkipTest("No test documents found in documents/ folder")

        # Find the Vendor Creation document
        cls.vendor_doc = None
        for doc in cls.test_docs:
            if "Vendor Creation" in doc.name:
                cls.vendor_doc = str(doc)
                break

    def test_01_parser_initialization(self):
        """Test parser initializes correctly"""
        self.assertIsNotNone(self.parser)
        self.assertIsInstance(self.parser, DocumentParser)

    def test_02_parse_all_documents(self):
        """Test that all documents can be parsed without errors"""
        for doc_path in self.test_docs:
            with self.subTest(document=doc_path.name):
                try:
                    parsed = self.parser.parse(str(doc_path))
                    self.assertIsNotNone(parsed)
                    self.assertEqual(parsed.file_path, str(doc_path))
                except Exception as e:
                    self.fail(f"Failed to parse {doc_path.name}: {str(e)}")

    def test_03_metadata_extraction(self):
        """Test metadata is extracted correctly"""
        for doc_path in self.test_docs:
            with self.subTest(document=doc_path.name):
                parsed = self.parser.parse(str(doc_path))

                # Metadata should exist
                self.assertIsNotNone(parsed.metadata)

                # At least some metadata should be present
                has_metadata = (
                    parsed.metadata.title or
                    parsed.metadata.author or
                    parsed.metadata.created
                )
                self.assertTrue(has_metadata, "No metadata extracted")

    def test_04_field_extraction_count(self):
        """Test that fields are extracted"""
        for doc_path in self.test_docs:
            with self.subTest(document=doc_path.name):
                parsed = self.parser.parse(str(doc_path))

                # Should extract some fields
                self.assertGreater(
                    len(parsed.all_fields),
                    0,
                    f"No fields extracted from {doc_path.name}"
                )

    def test_05_field_names_match_document(self):
        """Test that all field names in document are captured in JSON"""
        if not self.vendor_doc:
            self.skipTest("Vendor Creation document not found")

        # Extract field names from document
        doc_field_names = TextExtractor.extract_field_names_from_tables(self.vendor_doc)

        # Parse document
        parsed = self.parser.parse(self.vendor_doc)
        parsed_field_names = [f.name for f in parsed.all_fields]

        # Check coverage
        missing_fields = []
        for doc_field in doc_field_names:
            if doc_field not in parsed_field_names:
                missing_fields.append(doc_field)

        # Report findings
        coverage = (len(doc_field_names) - len(missing_fields)) / len(doc_field_names) * 100

        self.assertGreater(
            coverage,
            95,
            f"Field coverage is {coverage:.1f}%. Missing fields: {missing_fields[:5]}"
        )

    def test_06_field_attributes_completeness(self):
        """Test that field attributes are populated"""
        if not self.vendor_doc:
            self.skipTest("Vendor Creation document not found")

        parsed = self.parser.parse(self.vendor_doc)

        fields_with_names = sum(1 for f in parsed.all_fields if f.name)
        fields_with_types = sum(1 for f in parsed.all_fields if f.field_type != FieldType.UNKNOWN)
        fields_with_logic = sum(1 for f in parsed.all_fields if f.logic)

        # All fields should have names
        self.assertEqual(fields_with_names, len(parsed.all_fields))

        # Most fields should have types (>95%)
        type_coverage = fields_with_types / len(parsed.all_fields) * 100
        self.assertGreater(type_coverage, 95, f"Only {type_coverage:.1f}% fields have types")

        # Some fields should have logic
        self.assertGreater(fields_with_logic, 0, "No fields have logic/rules")

    def test_07_text_content_in_json(self):
        """Test that text content from document appears in JSON"""
        if not self.vendor_doc:
            self.skipTest("Vendor Creation document not found")

        # Extract text from document
        text_content = TextExtractor.extract_all_text(self.vendor_doc)

        # Parse document
        parsed = self.parser.parse(self.vendor_doc)
        json_str = json.dumps(parsed.to_dict()).lower()

        # Sample some field names and check they appear in JSON
        sample_fields = [
            "mobile number",
            "company code",
            "vendor",
            "created"
        ]

        missing_terms = []
        for term in sample_fields:
            if term.lower() not in json_str:
                missing_terms.append(term)

        self.assertEqual(
            len(missing_terms),
            0,
            f"Terms not found in JSON: {missing_terms}"
        )

    def test_08_workflow_extraction(self):
        """Test that workflows are extracted"""
        for doc_path in self.test_docs:
            with self.subTest(document=doc_path.name):
                parsed = self.parser.parse(str(doc_path))

                # Most documents should have workflows
                total_steps = sum(len(steps) for steps in parsed.workflows.values())

                if "Vendor Creation" in doc_path.name or "KYC" in doc_path.name:
                    self.assertGreater(
                        total_steps,
                        0,
                        f"No workflows extracted from {doc_path.name}"
                    )

    def test_09_workflow_actors_identified(self):
        """Test that workflow actors are correctly identified"""
        if not self.vendor_doc:
            self.skipTest("Vendor Creation document not found")

        parsed = self.parser.parse(self.vendor_doc)

        # Should have at least 2 different actors
        self.assertGreaterEqual(
            len(parsed.workflows),
            2,
            "Expected multiple workflow actors"
        )

        # Common actors should be present
        actors = set(parsed.workflows.keys())
        expected_actors = {'initiator', 'approver', 'spoc'}

        intersection = actors & expected_actors
        self.assertGreater(
            len(intersection),
            0,
            f"No expected actors found. Found: {actors}"
        )

    def test_10_mandatory_fields_detected(self):
        """Test that mandatory fields are correctly identified"""
        for doc_path in self.test_docs:
            with self.subTest(document=doc_path.name):
                parsed = self.parser.parse(str(doc_path))

                # Count mandatory fields
                mandatory_count = sum(1 for f in parsed.all_fields if f.is_mandatory)

                # Should have some mandatory fields
                if len(parsed.all_fields) > 10:
                    self.assertGreater(
                        mandatory_count,
                        0,
                        f"No mandatory fields detected in {doc_path.name}"
                    )

    def test_11_field_types_valid(self):
        """Test that field types are valid enum values"""
        for doc_path in self.test_docs:
            with self.subTest(document=doc_path.name):
                parsed = self.parser.parse(str(doc_path))

                for field in parsed.all_fields:
                    # Field type should be a valid FieldType enum
                    self.assertIsInstance(field.field_type, FieldType)

                    # Raw type should be preserved
                    self.assertIsInstance(field.field_type_raw, str)

    def test_12_dropdown_values_extracted(self):
        """Test that dropdown values are extracted"""
        if not self.vendor_doc:
            self.skipTest("Vendor Creation document not found")

        parsed = self.parser.parse(self.vendor_doc)

        # Should have dropdown mappings
        self.assertGreater(
            len(parsed.dropdown_mappings),
            0,
            "No dropdown mappings extracted"
        )

        # Check at least one dropdown has values
        has_values = any(len(values) > 0 for values in parsed.dropdown_mappings.values())
        self.assertTrue(has_values, "Dropdown mappings have no values")

    def test_13_table_extraction(self):
        """Test that tables are extracted"""
        for doc_path in self.test_docs:
            with self.subTest(document=doc_path.name):
                # Count tables in document
                doc = Document(str(doc_path))
                doc_table_count = len(doc.tables)

                if doc_table_count > 0:
                    parsed = self.parser.parse(str(doc_path))

                    # Should extract some table data (fields come from tables)
                    self.assertGreater(
                        len(parsed.all_fields),
                        0,
                        "Tables present but no fields extracted"
                    )

    def test_14_version_history_extraction(self):
        """Test version history extraction"""
        if not self.vendor_doc:
            self.skipTest("Vendor Creation document not found")

        parsed = self.parser.parse(self.vendor_doc)

        # Vendor Creation should have version history
        if len(parsed.version_history) > 0:
            entry = parsed.version_history[0]

            # Check required fields
            self.assertTrue(entry.version, "Version number missing")
            self.assertIsInstance(entry.version, str)

    def test_15_json_serialization(self):
        """Test that parsed data can be serialized to JSON"""
        for doc_path in self.test_docs:
            with self.subTest(document=doc_path.name):
                parsed = self.parser.parse(str(doc_path))

                try:
                    json_str = json.dumps(parsed.to_dict(), indent=2)
                    self.assertIsInstance(json_str, str)
                    self.assertGreater(len(json_str), 100)

                    # Should be valid JSON
                    data = json.loads(json_str)
                    self.assertIsInstance(data, dict)

                except Exception as e:
                    self.fail(f"Failed to serialize to JSON: {str(e)}")

    def test_16_vendor_creation_comprehensive(self):
        """Comprehensive test for Vendor Creation template"""
        if not self.vendor_doc:
            self.skipTest("Vendor Creation document not found")

        parsed = self.parser.parse(self.vendor_doc)

        # Test expected structure (updated for consecutive duplicate merging)
        # Note: Only consecutive duplicates are merged, not all duplicates across document
        assertions = [
            (len(parsed.all_fields) >= 300, "Should have 300+ fields"),
            (len(parsed.workflows) >= 2, "Should have multiple workflows"),
            (len(parsed.dropdown_mappings) >= 30, "Should have 30+ dropdown mappings"),
            (len(parsed.initiator_fields) > 0, "Should have initiator fields"),
            (len(parsed.spoc_fields) > 0, "Should have SPOC fields"),
            (len(parsed.approver_fields) > 0, "Should have approver fields"),
        ]

        failures = []
        for condition, message in assertions:
            if not condition:
                failures.append(message)

        self.assertEqual(len(failures), 0, f"Failures: {'; '.join(failures)}")

    def test_17_text_coverage_analysis(self):
        """Analyze what percentage of document text is captured in JSON"""
        if not self.vendor_doc:
            self.skipTest("Vendor Creation document not found")

        # Extract all text
        text_content = TextExtractor.extract_all_text(self.vendor_doc)

        # Parse document
        parsed = self.parser.parse(self.vendor_doc)
        json_str = json.dumps(parsed.to_dict()).lower()

        # Sample 50 random text snippets and check coverage
        import random
        text_snippets = [
            snippet.lower()
            for snippet in text_content['table_cells']
            if len(snippet) > 5 and len(snippet) < 50
        ]

        if len(text_snippets) > 50:
            text_snippets = random.sample(text_snippets, 50)

        found_count = sum(1 for snippet in text_snippets if snippet in json_str)
        coverage = (found_count / len(text_snippets)) * 100 if text_snippets else 0

        self.assertGreater(
            coverage,
            70,
            f"Only {coverage:.1f}% of text snippets found in JSON"
        )

    def test_18_field_logic_preservation(self):
        """Test that field logic and rules are preserved"""
        if not self.vendor_doc:
            self.skipTest("Vendor Creation document not found")

        parsed = self.parser.parse(self.vendor_doc)

        # Count fields with logic
        fields_with_logic = [f for f in parsed.all_fields if f.logic]

        self.assertGreater(
            len(fields_with_logic),
            100,
            "Expected most fields to have logic/rules"
        )

        # Check logic content quality
        substantial_logic = [
            f for f in fields_with_logic
            if len(f.logic) > 10  # More than just "Yes" or "No"
        ]

        self.assertGreater(
            len(substantial_logic),
            80,
            "Expected many fields with substantial logic"
        )

        # Check for merged logic (fields with multiple rules combined)
        # Note: Only consecutive duplicates are merged now
        merged_logic = [f for f in parsed.all_fields if '•' in f.logic]
        self.assertGreater(
            len(merged_logic),
            3,
            "Expected some fields with merged logic from consecutive duplicates"
        )

    def test_19_section_hierarchy(self):
        """Test that section hierarchy is preserved"""
        for doc_path in self.test_docs:
            with self.subTest(document=doc_path.name):
                parsed = self.parser.parse(str(doc_path))

                # Should have sections
                if len(parsed.sections) > 0:
                    # Check sections have proper structure
                    for section in parsed.sections:
                        self.assertIsInstance(section.heading, str)
                        self.assertIsInstance(section.level, int)
                        self.assertGreaterEqual(section.level, 1)

    def test_20_data_completeness_report(self):
        """Generate comprehensive completeness report"""
        if not self.vendor_doc:
            self.skipTest("Vendor Creation document not found")

        parsed = self.parser.parse(self.vendor_doc)

        report = {
            "total_fields": len(parsed.all_fields),
            "fields_with_types": sum(1 for f in parsed.all_fields if f.field_type != FieldType.UNKNOWN),
            "fields_with_logic": sum(1 for f in parsed.all_fields if f.logic),
            "mandatory_fields": sum(1 for f in parsed.all_fields if f.is_mandatory),
            "workflow_steps": sum(len(steps) for steps in parsed.workflows.values()),
            "dropdown_mappings": len(parsed.dropdown_mappings),
            "version_history": len(parsed.version_history),
        }

        # All metrics should be positive
        for metric, value in report.items():
            with self.subTest(metric=metric):
                self.assertGreater(value, 0, f"{metric} is zero")

        # Print report
        print("\n" + "="*80)
        print("DATA COMPLETENESS REPORT")
        print("="*80)
        for metric, value in report.items():
            print(f"  {metric:30s} : {value}")
        print("="*80)


class TestTextComparison(unittest.TestCase):
    """Test suite for comparing text extraction with JSON output"""

    @classmethod
    def setUpClass(cls):
        """Setup test environment"""
        cls.parser = DocumentParser()
        cls.vendor_doc = None

        # Find Vendor Creation document
        for doc_path in Path("documents").glob("*.docx"):
            if "Vendor Creation" in doc_path.name:
                cls.vendor_doc = str(doc_path)
                break

        if not cls.vendor_doc:
            raise unittest.SkipTest("Vendor Creation document not found")

    def test_all_field_names_present(self):
        """Verify all field names from document are in JSON"""
        # Extract from document
        doc_field_names = TextExtractor.extract_field_names_from_tables(self.vendor_doc)

        # Parse to JSON
        parsed = self.parser.parse(self.vendor_doc)
        json_field_names = [f.name for f in parsed.all_fields]

        # Compare
        missing = [name for name in doc_field_names if name not in json_field_names]

        self.assertEqual(
            len(missing),
            0,
            f"Missing {len(missing)} fields: {missing[:10]}"
        )

    def test_table_text_in_json(self):
        """Verify table content appears in JSON"""
        text_content = TextExtractor.extract_all_text(self.vendor_doc)
        parsed = self.parser.parse(self.vendor_doc)

        json_str = json.dumps(parsed.to_dict()).lower()

        # Sample table cells and verify presence
        sample_size = min(100, len(text_content['table_cells']))
        sample_cells = text_content['table_cells'][:sample_size]

        found = 0
        for cell_text in sample_cells:
            if len(cell_text) > 3 and cell_text.lower() in json_str:
                found += 1

        coverage = (found / sample_size) * 100
        self.assertGreater(coverage, 60, f"Only {coverage:.1f}% table text found in JSON")

    def test_no_data_loss(self):
        """Verify no significant data loss during parsing"""
        # Get document statistics
        doc = Document(self.vendor_doc)
        doc_table_count = len(doc.tables)
        doc_para_count = len([p for p in doc.paragraphs if p.text.strip()])

        # Get parsed statistics
        parsed = self.parser.parse(self.vendor_doc)

        # We should have substantial data
        assertions = [
            len(parsed.all_fields) > 100,
            len(parsed.workflows) > 0,
            sum(len(steps) for steps in parsed.workflows.values()) > 10,
        ]

        self.assertTrue(all(assertions), "Significant data appears to be missing")


def run_tests():
    """Run all tests and generate report"""
    # Create test suite
    loader = unittest.TestLoader()
    suite = unittest.TestSuite()

    # Add all test cases
    suite.addTests(loader.loadTestsFromTestCase(TestDocumentParser))
    suite.addTests(loader.loadTestsFromTestCase(TestTextComparison))

    # Run tests with verbose output
    runner = unittest.TextTestRunner(verbosity=2)
    result = runner.run(suite)

    # Print summary
    print("\n" + "="*80)
    print("TEST SUMMARY")
    print("="*80)
    print(f"Tests run:     {result.testsRun}")
    print(f"Successes:     {result.testsRun - len(result.failures) - len(result.errors)}")
    print(f"Failures:      {len(result.failures)}")
    print(f"Errors:        {len(result.errors)}")
    print(f"Skipped:       {len(result.skipped)}")
    print("="*80)

    return result.wasSuccessful()


if __name__ == "__main__":
    success = run_tests()
    exit(0 if success else 1)
