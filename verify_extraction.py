"""
Verification script to spot-check extraction quality against source documents.
"""

import json
from pathlib import Path
from docx import Document


def verify_vendor_creation_document():
    """
    Manually verify the Vendor Creation document extraction
    by comparing against actual document content.
    """

    doc_path = "documents/Vendor Creation Sample BUD(1).docx"
    result_path = "experiment_results/Vendor Creation Sample BUD(1)_full.json"

    print("="*100)
    print("VERIFICATION: Vendor Creation Sample BUD(1).docx")
    print("="*100)

    # Load the actual document
    doc = Document(doc_path)

    # Load the parsed results
    with open(result_path, 'r') as f:
        parsed = json.load(f)

    print("\n1. CHECKING TABLES IN DOCUMENT")
    print("-" * 100)

    table_count = len(doc.tables)
    # Count all table-related structures in parsed data
    parsed_tables = (
        len(parsed.get('reference_tables', [])) +
        sum(len(section.get('tables', [])) for section in parsed.get('sections', []))
    )

    print(f"Actual tables in document: {table_count}")
    print(f"Reference tables parsed: {len(parsed.get('reference_tables', []))}")
    print(f"Tables in sections: {sum(len(section.get('tables', [])) for section in parsed.get('sections', []))}")

    # Check first table structure
    print("\nFirst Table Verification:")
    if doc.tables:
        first_table = doc.tables[0]
        print(f"  Rows: {len(first_table.rows)}")
        print(f"  Columns: {len(first_table.columns) if first_table.rows else 0}")

        # Get headers
        if first_table.rows:
            headers = [cell.text.strip() for cell in first_table.rows[0].cells]
            print(f"  Headers: {headers[:5]}")  # First 5 headers only

    print("\n2. CHECKING FIELD EXTRACTION")
    print("-" * 100)

    # Find the main field definition table
    field_table = None
    for i, table in enumerate(doc.tables):
        if table.rows:
            headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
            if 'field name' in headers and 'field type' in headers:
                field_table = table
                print(f"Found field definition table at index {i}")
                break

    if field_table:
        # Count fields in actual table
        actual_field_count = len(field_table.rows) - 1  # Minus header
        parsed_field_count = len(parsed['all_fields'])

        print(f"Actual field rows in main table: {actual_field_count}")
        print(f"Total parsed fields: {parsed_field_count}")

        print("\nSample fields from document vs parsed:")
        print("-" * 100)

        # Check first 5 fields
        for i in range(1, min(6, len(field_table.rows))):
            row = field_table.rows[i]
            cells = [cell.text.strip() for cell in row.cells]

            if len(cells) >= 2:
                actual_name = cells[0]
                actual_type = cells[1]

                # Find in parsed
                parsed_field = None
                for field in parsed['all_fields']:
                    if field['name'] == actual_name:
                        parsed_field = field
                        break

                print(f"\nField {i}:")
                print(f"  Document: {actual_name} | {actual_type}")

                if parsed_field:
                    print(f"  Parsed:   {parsed_field['name']} | {parsed_field['field_type']} (raw: {parsed_field['field_type_raw']})")
                    if actual_name == parsed_field['name']:
                        print("  ✓ Name matches")
                    if parsed_field['field_type_raw'] == actual_type:
                        print("  ✓ Type matches")
                    else:
                        print(f"  ⚠ Type mismatch: '{actual_type}' != '{parsed_field['field_type_raw']}'")
                else:
                    print("  ✗ Field not found in parsed results")

    print("\n3. CHECKING FIELD TYPES DISTRIBUTION")
    print("-" * 100)

    field_types = {}
    for field in parsed['all_fields']:
        field_type = field['field_type']
        field_types[field_type] = field_types.get(field_type, 0) + 1

    for field_type, count in sorted(field_types.items(), key=lambda x: -x[1]):
        percentage = (count / len(parsed['all_fields'])) * 100
        print(f"  {field_type:20s} : {count:3d} ({percentage:5.1f}%)")

    print("\n4. CHECKING MANDATORY FIELDS")
    print("-" * 100)

    mandatory_count = sum(1 for f in parsed['all_fields'] if f['is_mandatory'])
    optional_count = len(parsed['all_fields']) - mandatory_count

    print(f"  Mandatory: {mandatory_count} ({(mandatory_count/len(parsed['all_fields'])*100):.1f}%)")
    print(f"  Optional:  {optional_count} ({(optional_count/len(parsed['all_fields'])*100):.1f}%)")

    # Sample mandatory fields
    print("\n  Sample mandatory fields:")
    mandatory_fields = [f for f in parsed['all_fields'] if f['is_mandatory']]
    for field in mandatory_fields[:5]:
        print(f"    - {field['name']} ({field['field_type']})")

    print("\n5. CHECKING WORKFLOWS")
    print("-" * 100)

    total_workflow_steps = sum(len(steps) for steps in parsed['workflows'].values())
    print(f"Total workflow steps extracted: {total_workflow_steps}")

    for actor, steps in parsed['workflows'].items():
        print(f"\n  {actor.upper()}: {len(steps)} steps")
        # Show first 2 steps
        for step in steps[:2]:
            desc = step['description'][:80] + "..." if len(step['description']) > 80 else step['description']
            action = f" [{step['action_type']}]" if step['action_type'] else ""
            print(f"    {step['step_number']}. {desc}{action}")

    print("\n6. CHECKING DROPDOWN MAPPINGS")
    print("-" * 100)

    dropdown_count = len(parsed['dropdown_mappings'])
    print(f"Total dropdown fields with values: {dropdown_count}")

    print("\nSample dropdown mappings:")
    for field_name, values in list(parsed['dropdown_mappings'].items())[:5]:
        print(f"\n  {field_name}:")
        print(f"    Values: {values[:5]}")

    print("\n7. CHECKING DOCUMENT REQUIREMENTS MATRIX")
    print("-" * 100)

    if parsed['document_requirements']:
        print(f"Document requirement entries: {len(parsed['document_requirements'])}")
        print("\nSample requirements:")
        for req in parsed['document_requirements'][:3]:
            print(f"\n  Document: {req['document_name']}")
            print(f"  Requirements: {req['requirements']}")
    else:
        print("No document requirements found")

    print("\n" + "="*100)
    print("VERIFICATION COMPLETE")
    print("="*100)

    # Overall assessment
    print("\nOVERALL ASSESSMENT:")

    checks = []
    checks.append(("Tables found", table_count > 0))
    checks.append(("Fields extracted", len(parsed['all_fields']) > 0))
    checks.append(("Field types identified", len(field_types) > 5))
    checks.append(("Workflows extracted", total_workflow_steps > 0))
    checks.append(("Dropdown values extracted", dropdown_count > 0))
    checks.append(("Mandatory flags set", mandatory_count > 0))

    passed = sum(1 for _, status in checks if status)
    total = len(checks)

    for check_name, status in checks:
        symbol = "✓" if status else "✗"
        print(f"  {symbol} {check_name}")

    print(f"\nScore: {passed}/{total} checks passed ({(passed/total*100):.0f}%)")

    if passed == total:
        print("\n🎉 ALL CHECKS PASSED - Document parsed correctly!")
    elif passed >= total * 0.8:
        print("\n✓ MOSTLY CORRECT - Minor issues to address")
    else:
        print("\n⚠ NEEDS ATTENTION - Significant issues found")


if __name__ == "__main__":
    verify_vendor_creation_document()
